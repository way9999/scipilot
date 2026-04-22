from __future__ import annotations

from datetime import datetime, timezone
import json
import re
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape
from zipfile import ZIP_DEFLATED, ZipFile
try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None  # type: ignore[assignment,misc]

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor, Emu
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[assignment]
    WD_SECTION = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    OxmlElement = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    Cm = None  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]

try:
    from lxml import etree
except ImportError:
    etree = None  # type: ignore[assignment]


ARTIFACT_SOURCES = {
    "paper": Path("drafts") / "paper-draft.md",
    "proposal": Path("drafts") / "proposal-draft.md",
    "literature_review": Path("drafts") / "literature-review.md",
    "research_answer": Path("drafts") / "research-answer.md",
    "presentation": Path("output") / "research-presentation.md",
}

PRESENTATION_JSON_PATH = Path("output") / "research-presentation.json"


def _now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _write_json(path: Path, payload: dict[str, Any]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return path


def _resolve_source(project_root: Path, artifact: str, source: str | Path | None) -> Path:
    if source:
        candidate = Path(source)
        return candidate if candidate.is_absolute() else (project_root / candidate).resolve()
    default = ARTIFACT_SOURCES.get(artifact)
    if not default:
        raise ValueError(f"Unsupported export artifact: {artifact}")
    return (project_root / default).resolve()


def _default_output_path(project_root: Path, source_path: Path, suffix: str) -> Path:
    safe_name = re.sub(r"[^0-9A-Za-z._-]+", "-", source_path.stem).strip("-") or "artifact"
    return project_root / "output" / "exports" / f"{safe_name}{suffix}"


def _markdown_lines(source_path: Path) -> list[str]:
    if not source_path.exists():
        raise FileNotFoundError(f"Export source not found: {source_path}")
    content = _repair_formula_text_for_export(source_path.read_text(encoding="utf-8"))
    return _repair_formula_text_for_export(content).splitlines()


def _split_markdown_table_row(line: str) -> list[str]:
    parts = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return [cell for cell in parts]


def _table_divider(line: str) -> bool:
    cells = _split_markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _is_table_caption(text: str) -> bool:
    """Detect table caption lines like '表4-1 ...' or 'Table 4-1 ...'."""
    return bool(re.match(r"^(?:表|Table)\s*\d+[-．.]\s*\d+", text.strip()))


def _is_figure_caption(text: str) -> bool:
    """Detect figure caption lines like '图4-1 ...' only (not English 'Figure X-Y').

    English 'Figure X-Y' lines in Chinese papers are typically redundant
    bilingual captions that duplicate the Chinese '图X-Y' line above.
    They are filtered out in the block parser instead.

    A real caption is a standalone description line, typically short,
    ending with a period or containing bilingual text (Chinese + English).
    Body text references contain verbs like '表明/显示/可知' after the figure number.
    """
    stripped = text.strip()
    # Only match Chinese figure captions, not English "Figure X-Y"
    if not re.match(r"^图\s*\d+[-．.]\s*\d+", stripped):
        return False
    # Body text references typically contain verbs indicating analysis
    # Only check AFTER the figure number prefix to avoid false positives
    # (e.g. "雷达图可以看出" should not trigger on "图")
    fig_prefix = re.match(r"^(?:图|Figure)\s*\d+[-．.]\s*\d+\s*", stripped)
    if fig_prefix:
        after_prefix = stripped[fig_prefix.end():]
        body_patterns = [
            r"表明", r"显示了", r"可以看出", r"可知", r"说明了", r"给出了",
            r"展示了", r"反映了", r"呈现出", r"可以看到", r"可以发现",
        ]
        for pat in body_patterns:
            if re.search(pat, after_prefix):
                return False
    # Captions are typically short (< 100 chars)
    if len(stripped) > 100:
        return False
    return True


def _is_foreign_caption(text: str) -> bool:
    """Detect English figure/table captions in a Chinese paper that duplicate the Chinese caption.

    Lines like 'Figure 2-1 ...' or 'Table 4-1 ...' are bilingual duplicates
    of the preceding '图2-1 ...' or '表4-1 ...' line and should be suppressed.
    """
    return bool(re.match(r"^(?:Figure|Table)\s+\d+[-．.]\s*\d+", text.strip()))


def _is_reference_line(text: str) -> bool:
    """Detect reference lines starting with [J], [C], [M], [D], [P], etc."""
    stripped = text.strip()
    if not stripped:
        return False
    # Match numbered references like "1. Author..." with academic indicators
    # Must contain typical reference markers: [J], [C], [M], DOI, year, etc.
    if re.match(r"^\d+\s*[.．]\s*\S", stripped):
        has_ref_marker = bool(re.search(
            r"\[?[JCMDCP]\]?\s*[.:：]|DOI[:\s]|http|ISBN|ISBN|Vol\.|Ed\.",
            stripped, re.IGNORECASE
        ))
        if has_ref_marker:
            return True
        # Also match if it has author-like pattern + year
        if re.match(r"^\d+\s*[.．]\s*[A-Z][a-z]+(\s+[A-Z])?\s*,", stripped):
            return True
    if re.match(r"^\[\d+\]", stripped):
        return True
    return False


def _normalize_equation_tag(tag: str | None) -> str | None:
    raw = str(tag or "").strip()
    if not raw:
        return None
    raw = re.sub(r"^[（(]\s*", "", raw)
    raw = re.sub(r"[）)]$", "", raw)
    raw = re.sub(r"^\s*式\s*", "", raw)
    return raw.strip() or None


def _display_equation_tag_text(tag: str | None) -> str | None:
    normalized = _normalize_equation_tag(tag)
    if not normalized:
        return None
    return f"（式{normalized}）"


def _looks_like_equation_line(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped:
        return False
    narrative = stripped.lstrip("。；，,.:：!?！？ ")
    if narrative.startswith(("将式", "由式", "如式", "where", "thus", "therefore", "hence")):
        return False
    if stripped in {"$$", "$"}:
        return True
    if stripped.startswith(("#", "!", "|", ">", "```", "- ", "* ")):
        return False
    if stripped.startswith(("图", "Figure", "表", "Table")):
        return False
    math_hints = ("=", "\\", "^", "_", "{", "}", "[", "]", "(", ")", "sin", "cos", "log", "exp", "sqrt", "sum", "int", "theta", "lambda")
    if any(token in stripped for token in math_hints):
        return True
    return bool(
        re.fullmatch(r"[A-Za-z0-9\s\+\-\*/\.,:&<>|]+", stripped)
        and any(ch.isalpha() for ch in stripped)
        and any(op in stripped for op in "+-*/")
    )


def _looks_like_explicit_equation_body(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped:
        return False
    narrative = stripped.lstrip("。；，,.:：!?！？ ")
    if narrative.startswith(("将式", "由式", "如式", "where", "thus", "therefore", "hence")):
        return False
    if stripped.startswith(("#", "!", "|", ">", "```", "- ", "* ", "图", "Figure", "表", "Table")):
        return False
    strong_tokens = ("=", "\\frac", "\\sum", "\\int", "\\begin{", "\\rightarrow")
    if any(token in stripped for token in strong_tokens):
        return True
    return bool(
        len(stripped) <= 120
        and re.fullmatch(r"[A-Za-z0-9\\\s\+\-\*/\^\_\(\)\[\]\{\}\.,:&<>|]+", stripped)
        and any(token in stripped for token in ("\\", "^", "_"))
    )


def _repair_formula_text_for_export(text: str) -> str:
    repaired = str(text or "")
    if not repaired:
        return repaired

    repaired = repaired.replace("\r\n", "\n")
    repaired = re.sub(r"(?m)^\s*\\tag\{([^}]+)\}\s*\$\$\s*$", lambda m: f"\\tag{{{m.group(1).strip()}}}\n$$", repaired)
    repaired = re.sub(r"(?m)^\s*\$\$\s*\\tag\{([^}]+)\}\s*$", lambda m: f"$$\n\\tag{{{m.group(1).strip()}}}", repaired)

    repaired = re.sub(
        r"\$\$\s*([\s\S]*?)\s*\\tag\{([^}]+)\}\s*([。；;，,.])(?=[^\n$])",
        lambda m: "$$\n"
        + m.group(1).strip()
        + f"\n\\tag{{{m.group(2).strip()}}}\n$$"
        + m.group(3),
        repaired,
    )
    repaired = re.sub(
        r"\$\$\s*([\s\S]*?)\s*\\tag\{([^}]+)\}\s*(?=\n(?:将式|由式|如式|where|thus|therefore|hence)\b)",
        lambda m: "$$\n"
        + m.group(1).strip()
        + f"\n\\tag{{{m.group(2).strip()}}}\n$$\n",
        repaired,
        flags=re.I,
    )
    repaired = re.sub(
        r"\$\$\s*([\s\S]*?)\n((?:将式|由式|如式|where|thus|therefore|hence)[^\n]*)\n\\tag\{([^}]+)\}\n\$\$",
        lambda m: "$$\n"
        + m.group(1).strip()
        + f"\n\\tag{{{m.group(3).strip()}}}\n$$\n"
        + m.group(2).strip(),
        repaired,
        flags=re.I,
    )

    if repaired.count("$$") % 2 == 1 and "\\tag{" in repaired:
        repaired = re.sub(
            r"\$\$\s*([\s\S]*?)\s*\\tag\{([^}]+)\}",
            lambda m: "$$\n"
            + m.group(1).strip()
            + f"\n\\tag{{{m.group(2).strip()}}}\n$$",
            repaired,
            count=1,
        )

    return repaired


def _parse_blockquote_image_placeholder(block_lines: list[str]) -> dict[str, Any] | None:
    normalized_lines = [re.sub(r"^\s*>\s?", "", line).strip() for line in block_lines if line.strip()]
    if not normalized_lines:
        return None

    header = normalized_lines[0]
    match = re.match(
        r"^\[(?:(?:此处插入|请插入|insert)\s*(?:图|Figure|Fig\.?)\s*(\d+[-．.]\d+)|待补图|Figure Placeholder)\]\s*(.*)$",
        header,
        re.IGNORECASE,
    )
    if not match:
        return None

    ref = (match.group(1) or "").replace("．", "-").replace(".", "-")
    caption = match.group(2).strip()
    payload: dict[str, Any] = {"type": "image_placeholder", "ref": ref, "caption": caption}

    for extra in normalized_lines[1:]:
        key, sep, value = extra.partition("：")
        if not sep:
            key, sep, value = extra.partition(":")
        normalized_key = key.strip().lower()
        normalized_value = value.strip() if sep else extra.strip()
        if normalized_key in {"图型建议", "图示建议", "type"}:
            payload["figure_type"] = normalized_value
        elif normalized_key in {"应展示内容", "展示内容", "required content"}:
            payload["goal"] = normalized_value
        elif normalized_key in {"推荐素材来源", "素材来源", "suggested source"}:
            payload["evidence"] = normalized_value
        elif normalized_value and not payload.get("goal"):
            payload["goal"] = normalized_value

    return payload


def _rescue_formula_segments(text: str) -> list[dict[str, Any]]:
    lines = [line.strip() for line in _repair_formula_text_for_export(text).splitlines() if line.strip()]
    if not lines:
        return []

    body_lines: list[str] = []
    trailing_lines: list[str] = []
    tag: str | None = None
    in_formula = False

    for line in lines:
        if line == "$$":
            in_formula = True
            continue
        tag_matches = re.findall(r"\\tag\{([^}]+)\}", line)
        if tag_matches and tag is None:
            tag = _normalize_equation_tag(tag_matches[0])
        cleaned = re.sub(r"\s*\\tag\{[^}]+\}\s*", " ", line).strip().strip("$").strip()
        if not cleaned:
            continue
        if in_formula and not _looks_like_explicit_equation_body(cleaned) and not _looks_like_equation_line(cleaned):
            trailing_lines.append(cleaned)
            continue
        if _looks_like_explicit_equation_body(cleaned) or _looks_like_equation_line(cleaned):
            in_formula = True
            body_lines.append(cleaned)
        elif in_formula:
            trailing_lines.append(cleaned)

    if not body_lines:
        return []

    rescued: list[dict[str, Any]] = [{"type": "display_formula", "content": "\n".join(body_lines).strip(), "tag": tag}]
    if trailing_lines:
        rescued.append({"type": "text", "content": "\n".join(trailing_lines).strip()})
    return rescued


def _extract_formulas_from_text(text: str) -> list[dict[str, Any]]:
    """Extract $$...$$ and $...$ from paragraph text, returning segments.

    Returns list of segments: {"type": "text"/"display_formula"/"inline_formula", "content": "..."}
    """
    text = _repair_formula_text_for_export(text)
    segments: list[dict[str, Any]] = []
    # First extract display formulas $$...$$ (greedy to capture full formula)
    display_pattern = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
    last_end = 0
    for m in display_pattern.finditer(text):
        if m.start() > last_end:
            before = text[last_end:m.start()]
            if before.strip():
                segments.append({"type": "text", "content": before})
        formula = m.group(1).strip()
        # Extract \tag{X.Y}
        tag_matches = re.findall(r"\\tag\{([^}]+)\}", formula)
        tag = _normalize_equation_tag(tag_matches[0] if tag_matches else None)
        if tag_matches:
            formula = re.sub(r"\s*\\tag\{[^}]+\}\s*", " ", formula).strip()
        if not _looks_like_explicit_equation_body(formula) and not _looks_like_equation_line(formula):
            if formula.strip():
                segments.append({"type": "text", "content": formula})
            last_end = m.end()
            continue
        segments.append({"type": "display_formula", "content": formula, "tag": tag})
        last_end = m.end()
    remaining = text[last_end:]
    if remaining.strip():
        # Extract inline formulas $...$ from remaining text
        inline_pattern = re.compile(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)")
        inline_last = 0
        for m2 in inline_pattern.finditer(remaining):
            if m2.start() > inline_last:
                txt = remaining[inline_last:m2.start()]
                if txt.strip():
                    segments.append({"type": "text", "content": txt})
            segments.append({"type": "inline_formula", "content": m2.group(1).strip()})
            inline_last = m2.end()
        if inline_last < len(remaining):
            txt = remaining[inline_last:]
            if txt.strip():
                segments.append({"type": "text", "content": txt})
    cleaned_segments: list[dict[str, Any]] = []
    for seg in segments:
        if seg["type"] != "text":
            cleaned_segments.append(seg)
            continue
        cleaned_text = re.sub(r"(?m)^\s*\\tag\{[^}]+\}\s*$", "", seg["content"])
        cleaned_text = re.sub(r"(?m)^\s*\$\$\s*$", "", cleaned_text)
        if cleaned_text.strip():
            cleaned_segments.append({"type": "text", "content": cleaned_text})
    segments = cleaned_segments
    if not any(seg["type"] in {"display_formula", "inline_formula"} for seg in segments) and ("$$" in text or "\\tag{" in text):
        rescued = _rescue_formula_segments(text)
        if rescued:
            return rescued
    return segments


def _iter_markdown_blocks(lines: list[str]) -> list[dict[str, Any]]:
    """Enhanced markdown block parser.

    Detects: headings, paragraphs, tables, code, images, formulas ($$...$$),
    bullets, numbered items, pagebreaks, TOC, figure captions, table captions,
    and reference lines.
    """
    blocks: list[dict[str, Any]] = []
    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []
    in_code = False
    code_buffer: list[str] = []

    # Lines ending with any of these are treated as complete logical units;
    # adjacent terminated lines in the same block are split into separate
    # paragraphs rather than joined with a soft break. Lines not ending with
    # a terminator are treated as soft-wrapped prose and joined normally.
    _sentence_terminators = ("。", "！", "？", "；", "：", ".", "!", "?")

    def _group_soft_lines(buf: list[str]) -> list[str]:
        groups: list[list[str]] = []
        current: list[str] = []
        for ln in buf:
            s = ln.rstrip()
            if not s:
                continue
            current.append(s)
            if s.endswith(_sentence_terminators):
                groups.append(current)
                current = []
        if current:
            groups.append(current)
        out: list[str] = []
        for g in groups:
            if len(g) == 1:
                out.append(g[0])
            else:
                # Soft-wrapped prose: join with space for ASCII, empty for CJK.
                joined = ""
                for i, piece in enumerate(g):
                    if i == 0:
                        joined = piece
                    else:
                        sep = " " if (joined[-1].isascii() and piece[0].isascii()) else ""
                        joined += sep + piece
                out.append(joined)
        return out

    def flush_paragraph() -> None:
        if paragraph_buffer:
            placeholder_block = _parse_blockquote_image_placeholder(paragraph_buffer)
            if placeholder_block:
                blocks.append(placeholder_block)
                paragraph_buffer.clear()
                return
            texts = _group_soft_lines(paragraph_buffer)
            for text in texts:
                text = text.strip()
                if not text:
                    continue
                # Extract inline image placeholders: [此处插入图X-Y]
                # Split paragraph into text + placeholder + text segments
                inline_pattern = re.compile(r'\[(?:此处插入|请插入|insert)\s*(?:图|Figure|Fig\.?)\s*(\d+[-．.]\d+)\]', re.IGNORECASE)
                parts = inline_pattern.split(text)
                # parts alternates: text, ref, text, ref, text, ...
                for idx, part in enumerate(parts):
                    if idx % 2 == 0:
                        # Text segment
                        part = part.strip()
                        if not part:
                            continue
                        # Check if this paragraph contains display formulas
                        if "$$" in part or "\\tag{" in part:
                            blocks.append({"type": "formula_paragraph", "text": part})
                        elif _is_table_caption(part):
                            blocks.append({"type": "table_caption", "text": part})
                        elif _is_figure_caption(part):
                            blocks.append({"type": "figure_caption", "text": part})
                        elif _is_reference_line(part):
                            blocks.append({"type": "reference", "text": part})
                        elif _is_foreign_caption(part):
                            pass  # Skip bilingual duplicate (Figure/Table in Chinese paper)
                        else:
                            blocks.append({"type": "paragraph", "text": part})
                    else:
                        # Figure reference (odd index)
                        blocks.append({"type": "image_placeholder", "ref": part.replace("．", "-").replace(".", "-")})
            paragraph_buffer.clear()

    def flush_table() -> None:
        if table_buffer:
            rows = [_split_markdown_table_row(row) for row in table_buffer
                    if row.strip() and not _table_divider(row)]
            if rows:
                blocks.append({"type": "table", "rows": rows})
            table_buffer.clear()

    def flush_code() -> None:
        if code_buffer:
            blocks.append({"type": "code", "text": "\n".join(code_buffer).rstrip()})
            code_buffer.clear()

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith(">"):
            if paragraph_buffer and not all(buf_line.lstrip().startswith(">") for buf_line in paragraph_buffer if buf_line.strip()):
                flush_paragraph()
            paragraph_buffer.append(stripped)
            continue
        if paragraph_buffer and all(buf_line.lstrip().startswith(">") for buf_line in paragraph_buffer if buf_line.strip()):
            flush_paragraph()

        if stripped.startswith("```"):
            flush_paragraph()
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        # ---- Table handling ----
        # We accumulate table lines into table_buffer and flush when
        # a non-table line is encountered (or at end of file).
        # Two table formats are supported:
        #   A) Standard multi-line markdown tables (rows on separate lines)
        #   B) Compressed single-line tables (entire table on one line)
        if stripped.startswith("|"):
            flush_paragraph()

            # Detect compressed table FIRST (before trailing text extraction),
            # because compressed tables may not end with "|" (LLM-generated
            # markdown sometimes omits the final pipe).
            full_divider = re.search(r"\|(\s*-{2,}[^|]*\|){2,}\s*-{2,}[^|]*\|", stripped)

            # A line starting with "|" is only a real table row if:
            # 1) It is a compressed table (has full_divider with content cells), OR
            # 2) It ends with "|" and has multiple pipe-separated cells (standard table row), OR
            # 3) We are already accumulating a table (table_buffer has content)
            is_compressed = False
            is_table_line = False
            if full_divider:
                before = stripped[:full_divider.start()].strip()
                if before.startswith("|") and before.endswith("|"):
                    inner = before.strip().strip("|").strip()
                    cells = [c.strip() for c in inner.split("|")]
                    content_cells = [c for c in cells if c and not re.fullmatch(r":?-{2,}:?", c)]
                    if len(content_cells) >= 2:
                        is_compressed = True
                        is_table_line = True
                elif table_buffer:
                    # Pure divider line inside a table (full_divider matched the
                    # entire line so before is empty).  Still a table row.
                    is_table_line = True
            elif stripped.endswith("|") and table_buffer:
                is_table_line = True
            elif stripped.endswith("|"):
                inner = stripped.strip().strip("|").strip()
                cells = [c.strip() for c in inner.split("|")]
                non_empty = [c for c in cells if c and not re.fullmatch(r":?-{2,}:?", c)]
                if len(non_empty) >= 2:
                    is_table_line = True

            if not is_table_line:
                paragraph_buffer.append(stripped)
                continue

            # For compressed tables, don't extract trailing text — the whole
            # line is table data. For standard table rows, extract trailing.
            trailing_text = ""
            if not is_compressed and not stripped.endswith("|"):
                last_pipe = stripped.rfind("|")
                if last_pipe > 0:
                    after_pipe = stripped[last_pipe + 1:].strip()
                    if after_pipe and not re.match(r"^[\s|:;-]+$", after_pipe):
                        trailing_text = after_pipe
                        stripped = stripped[:last_pipe + 1]

            if is_compressed:
                # Compressed single-line table — parse inline
                before = stripped[:full_divider.start()].strip()
                table_buffer.append(before)

                header_cols = len(_split_markdown_table_row(before))

                after = stripped[full_divider.end():].strip()
                if after.startswith("|"):
                    after_inner = after.strip().strip("|").strip()
                    cells = [c.strip() for c in after_inner.split("|")]
                    row_idx = 0
                    while row_idx < len(cells):
                        potential_row = cells[row_idx:row_idx + header_cols]
                        if len(potential_row) == header_cols:
                            if all(re.fullmatch(r":?-{2,}:?", c) for c in potential_row):
                                row_idx += header_cols
                                continue
                            table_buffer.append("| " + " | ".join(potential_row) + " |")
                        row_idx += header_cols
                # Flush immediately (all data is on one line)
                flush_table()
            else:
                # Standard multi-line table row (header, data, or divider)
                table_buffer.append(stripped)

            # Handle trailing content
            if trailing_text:
                # Generic image placeholder detection
                img_ph = re.search(
                    r"\[(?:此处插入|请插入|insert)\s*(?:图|Figure|Fig\.?)\s*(\d+[-．.]\d+)\]",
                    trailing_text, re.IGNORECASE,
                )
                if img_ph:
                    blocks.append({"type": "image_placeholder", "ref": img_ph.group(1)})
                elif trailing_text.strip():
                    blocks.append({"type": "paragraph", "text": trailing_text.strip()})
            continue

        # If we were accumulating a table and hit a non-table line, flush it
        if table_buffer and not stripped.startswith("|"):
            flush_table()

        if stripped == "[PAGEBREAK]":
            flush_paragraph()
            blocks.append({"type": "pagebreak"})
            continue

        if stripped == "[TOC]":
            flush_paragraph()
            blocks.append({"type": "toc"})
            continue

        # Generic image placeholder detection
        img_placeholder = re.match(
            r"^\[(?:此处插入|请插入|insert)\s*(?:图|Figure|Fig\.?)\s*(\d+[-．.]\d+)\]",
            stripped, re.IGNORECASE,
        )
        if img_placeholder:
            flush_paragraph()
            blocks.append({"type": "image_placeholder", "ref": img_placeholder.group(1)})
            continue

        # Skip standalone English "Figure/Table X-Y ..." lines in Chinese papers —
        # the corresponding "图X-Y"/"表X-Y" caption already exists above the image.
        if _is_foreign_caption(stripped) and len(stripped) < 200:
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.+?)\)$", stripped)
        if image_match:
            flush_paragraph()
            blocks.append(
                {
                    "type": "image",
                    "alt": image_match.group(1).strip(),
                    "path": image_match.group(2).strip(),
                }
            )
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            blocks.append({"type": "heading", "level": level, "text": heading_match.group(2).strip()})
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            flush_paragraph()
            blocks.append({"type": "bullet", "text": bullet_match.group(1).strip()})
            continue

        # Check reference BEFORE numbered (references also start with "N. ")
        if _is_reference_line(stripped):
            flush_paragraph()
            blocks.append({"type": "reference", "text": stripped})
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            flush_paragraph()
            blocks.append({"type": "numbered", "text": number_match.group(1).strip()})
            continue

        if not stripped:
            flush_paragraph()
            flush_table()
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph()
    flush_table()
    if in_code:
        flush_code()
    return blocks


def _convert_latex_to_omml(latex_str: str) -> str | None:
    """Convert a LaTeX formula to OMML XML for Word native equation rendering.

    Uses latex2mathml → MathML → OMML XSLT transform.
    Returns OMML XML string or None on failure.
    """
    if etree is None:
        return None
    try:
        from latex2mathml.converter import convert as l2m_convert
    except ImportError:
        return None

    try:
        # Pre-process common LaTeX commands that latex2mathml may not handle
        clean = latex_str.strip()
        # Replace \mathcal, \mathbb, etc. with simpler forms if needed
        clean = re.sub(r"\\mathcal\{(\w)\}", r"\\mathcal{\1}", clean)
        clean = re.sub(r"\\mathbb\{(\w)\}", r"\\mathbb{\1}", clean)
        # Remove trailing punctuation that's outside the formula
        clean = clean.rstrip("。.,;:：；，")

        mathml_str = l2m_convert(clean)

        # MathML → OMML via XSLT
        # The XSLT is embedded as a string (from Microsoft's MML2OMML.XSL)
        xslt_path = _find_mml2omml()
        if xslt_path and Path(xslt_path).exists():
            mathml_doc = etree.fromstring(mathml_str.encode("utf-8"))
            xslt = etree.parse(str(xslt_path))
            transform = etree.XSLT(xslt)
            omml_doc = transform(mathml_doc)
            omml_str = etree.tostring(omml_doc, encoding="unicode")
            return _postprocess_omml(omml_str)
        else:
            # Fallback: use the bundled XSLT
            xslt_str = _get_bundled_mml2omml_xslt()
            if xslt_str:
                mathml_doc = etree.fromstring(mathml_str.encode("utf-8"))
                xslt_doc = etree.fromstring(xslt_str.encode("utf-8"))
                transform = etree.XSLT(xslt_doc)
                omml_doc = transform(mathml_doc)
                omml_str = etree.tostring(omml_doc, encoding="unicode")
                return _postprocess_omml(omml_str)
    except Exception:
        pass
    return None


def _postprocess_omml(omml_str: str) -> str:
    """Post-process OMML XML to fix known issues from MML2OMML.XSL conversion.

    Fixes:
    - m:limUpp with single child → m:acc (accent)
    """
    if etree is None:
        return omml_str
    try:
        ns_m = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        root = etree.fromstring(omml_str.encode("utf-8"))
        fixed = 0
        for limUpp in root.iter(f"{{{ns_m}}}limUpp"):
            children = list(limUpp)
            # m:limUpp with only one child element → convert to m:acc
            elems = [c for c in children if c.tag.endswith("}e") or c.tag.endswith("}r")]
            if len(elems) == 1:
                acc = etree.Element(f"{{{ns_m}}}acc")
                accPr = etree.SubElement(acc, f"{{{ns_m}}}accPr")
                # Try to detect accent character from second child
                non_e = [c for c in children if c.tag.endswith("}lim")]
                if non_e:
                    lim_text_elts = non_e[0].findall(f".//{{{ns_m}}}t")
                    if lim_text_elts:
                        chr_val = lim_text_elts[0].text or "^"
                        chr_prop = etree.SubElement(accPr, f"{{{ns_m}}}chr")
                        chr_prop.set(f"{{{ns_m}}}val", chr_val)
                # Move the element from limUpp to acc
                parent = limUpp.getparent()
                idx = list(parent).index(limUpp)
                parent.remove(limUpp)
                acc.append(elems[0])
                parent.insert(idx, acc)
                fixed += 1
        if fixed:
            return etree.tostring(root, encoding="unicode")
    except Exception:
        pass
    return omml_str


def _find_mml2omml() -> str | None:
    """Find MML2OMML.XSL on the system (comes with Microsoft Office)."""
    import os

    # Common Office installation paths on Windows
    candidates = [
        os.path.expandvars(r"%ProgramFiles%\Microsoft Office\root\Office16\MML2OMML.XSL"),
        os.path.expandvars(r"%ProgramFiles(x86)%\Microsoft Office\root\Office16\MML2OMML.XSL"),
        os.path.expandvars(r"%ProgramFiles%\Microsoft Office\Office16\MML2OMML.XSL"),
        os.path.expandvars(r"%ProgramFiles(x86)%\Microsoft Office\Office16\MML2OMML.XSL"),
        os.path.expandvars(r"%ProgramFiles%\Microsoft Office\root\Office17\MML2OMML.XSL"),
        os.path.expandvars(r"%ProgramFiles%\Microsoft Office\root\Office19\MML2OMML.XSL"),
        os.path.expandvars(r"%ProgramFiles%\Microsoft Office\root\Office21\MML2OMML.XSL"),
        os.path.expandvars(r"%ProgramFiles%\Microsoft Office\root\Office24\MML2OMML.XSL"),
    ]
    for c in candidates:
        if os.path.isfile(c):
            return c
    return None


def _get_bundled_mml2omml_xslt() -> str | None:
    """Return a minimal MML2OMML XSLT for basic formula conversion.

    This is a simplified version that handles common LaTeX constructs.
    """
    return '''<?xml version="1.0" encoding="UTF-8"?>
<xsl:stylesheet version="1.0"
    xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
    xmlns:mml="http://www.w3.org/1998/Math/MathML"
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    exclude-result-prefixes="mml">

  <xsl:output method="xml" encoding="UTF-8"/>

  <xsl:template match="mml:math">
    <m:oMathPara><m:oMath>
      <xsl:apply-templates/>
    </m:oMath></m:oMathPara>
  </xsl:template>

  <xsl:template match="mml:mrow">
    <xsl:apply-templates/>
  </xsl:template>

  <xsl:template match="mml:mi|mml:mn|mml:mo|mml:mtext">
    <m:r>
      <m:t><xsl:value-of select="."/></m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:msup">
    <m:sSup>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:sup><xsl:apply-templates select="*[2]"/></m:sup>
    </m:sSup>
  </xsl:template>

  <xsl:template match="mml:msub">
    <m:sSub>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
    </m:sSub>
  </xsl:template>

  <xsl:template match="mml:msubsup">
    <m:sSubSup>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
      <m:sup><xsl:apply-templates select="*[3]"/></m:sup>
    </m:sSubSup>
  </xsl:template>

  <xsl:template match="mml:mfrac">
    <m:f>
      <m:num><xsl:apply-templates select="*[1]"/></m:num>
      <m:den><xsl:apply-templates select="*[2]"/></m:den>
    </m:f>
  </xsl:template>

  <xsl:template match="mml:mover">
    <m:acc>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
    </m:acc>
  </xsl:template>

  <xsl:template match="mml:munder">
    <m:limLow>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:lim><xsl:apply-templates select="*[2]"/></m:lim>
    </m:limLow>
  </xsl:template>

  <xsl:template match="mml:munderover">
    <m:nary>
      <m:naryPr>
        <m:chr m:val="∫"/>
      </m:naryPr>
      <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
      <m:sup><xsl:apply-templates select="*[3]"/></m:sup>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
    </m:nary>
  </xsl:template>

  <xsl:template match="mml:msqrt">
    <m:rad>
      <m:radPr><m:degHide m:val="1"/></m:radPr>
      <m:deg/>
      <m:e><xsl:apply-templates/></m:e>
    </m:rad>
  </xsl:template>

  <xsl:template match="mml:mroot">
    <m:rad>
      <m:radPr><m:degHide m:val="0"/></m:radPr>
      <m:deg><xsl:apply-templates select="*[2]"/></m:deg>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
    </m:rad>
  </xsl:template>

  <xsl:template match="mml:mfenced">
    <m:d>
      <m:dPr>
        <xsl:if test="@open">
          <m:begChr m:val="{@open}"/>
        </xsl:if>
        <xsl:if test="@close">
          <m:endChr m:val="{@close}"/>
        </xsl:if>
      </m:dPr>
      <m:e><xsl:apply-templates/></m:e>
    </m:d>
  </xsl:template>

  <xsl:template match="mml:mtable">
    <m:m>
      <xsl:apply-templates/>
    </m:m>
  </xsl:template>

  <xsl:template match="mml:mtr">
    <m:mr>
      <xsl:apply-templates/>
    </m:mr>
  </xsl:template>

  <xsl:template match="mml:mtd">
    <m:e>
      <xsl:apply-templates/>
    </m:e>
  </xsl:template>

  <xsl:template match="mml:mstyle|mml:mphantom">
    <xsl:apply-templates/>
  </xsl:template>

  <xsl:template match="text()">
    <!-- Skip whitespace-only text nodes -->
  </xsl:template>

</xsl:stylesheet>'''


def _insert_omml_formula(paragraph: Any, omml_xml: str) -> None:
    """Insert OMML XML into a python-docx paragraph."""
    if OxmlElement is None or paragraph is None:
        return
    try:
        omml_elem = etree.fromstring(omml_xml.encode("utf-8"))
        ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        local_tag = omml_elem.tag.split("}")[-1] if "}" in omml_elem.tag else omml_elem.tag

        if local_tag == "oMathPara":
            # Wrapper: extract child oMath elements
            for child in omml_elem:
                child_local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if child_local == "oMath":
                    paragraph._p.append(child)
        elif local_tag == "oMath":
            # Direct oMath element
            paragraph._p.append(omml_elem)
        else:
            # Unknown wrapper, try to find oMath descendants
            for elem in omml_elem.iter(f"{{{ns}}}oMath"):
                paragraph._p.append(elem)
    except Exception:
        # Fallback: insert as plain text
        run = paragraph.add_run(_strip_xml_tags(omml_xml))
        run.font.name = "Cambria Math"
        if qn is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")


def _strip_xml_tags(xml_str: str) -> str:
    return re.sub(r"<[^>]+>", "", xml_str)


def _build_three_line_table(doc: Any, headers: list[str], data_rows: list[list[str]],
                             style: str = "thesis") -> Any:
    """Build an academic three-line table (三线表).

    Top line: 1.5pt, header separator: 0.75pt, bottom line: 1.5pt.
    No vertical lines, no intermediate horizontal lines.
    Headers: bold, centered, 10.5pt SimSun/Times New Roman.
    Data: centered, 10.5pt.
    """
    if Document is None or OxmlElement is None or qn is None or Pt is None:
        # Fallback to basic table
        num_rows = len(data_rows) + 1
        num_cols = len(headers)
        return doc.add_table(rows=num_rows, cols=num_cols)

    num_rows = len(data_rows) + 1
    num_cols = len(headers)
    table = doc.add_table(rows=num_rows, cols=num_cols)

    # Remove default borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)
    # Remove existing borders
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

    # Set table width to full page
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblW)

    # Set table alignment center
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    existing_jc = tblPr.find(qn("w:jc"))
    if existing_jc is not None:
        tblPr.remove(existing_jc)
    tblPr.append(jc)

    # Top border: 1.5pt (sz=12)
    top_border = OxmlElement("w:top")
    top_border.set(qn("w:val"), "single")
    top_border.set(qn("w:sz"), "12")
    top_border.set(qn("w:space"), "0")
    top_border.set(qn("w:color"), "000000")
    borders.append(top_border)

    # Bottom border: 1.5pt (sz=12)
    bottom_border = OxmlElement("w:bottom")
    bottom_border.set(qn("w:val"), "single")
    bottom_border.set(qn("w:sz"), "12")
    bottom_border.set(qn("w:space"), "0")
    bottom_border.set(qn("w:color"), "000000")
    borders.append(bottom_border)

    # Fill header row
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        # Clear all existing runs (cell.text="" doesn't remove runs)
        for p in cell.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Headers usually don't contain inline math, but support it for safety.
        _apply_inline_formatting(p, header_text, style=style)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.bold = True
        # Top border (1.5pt) + bottom border (0.75pt) on header cells
        _set_cell_border(
            cell,
            top={"val": "single", "sz": "12", "color": "000000"},
            bottom={"val": "single", "sz": "6", "color": "000000"},
        )

    # Fill data rows
    for row_idx, row_data in enumerate(data_rows):
        is_last_row = (row_idx == len(data_rows) - 1)
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = table.cell(row_idx + 1, col_idx)
            for p in cell.paragraphs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Honour inline markdown ($math$, *italic*, **bold**) inside cells so
            # variables like $M$ or $c_{ij}$ render as italic math, not literal.
            _apply_inline_formatting(p, str(cell_text), style=style)
            for run in p.runs:
                if run.font.size is None:
                    run.font.size = Pt(10.5)
            # Bottom border (1.5pt) on last row
            if is_last_row:
                _set_cell_border(cell, bottom={"val": "single", "sz": "12", "color": "000000"})

    return table


def _build_table(doc: Any, headers: list[str], data_rows: list[list[str]],
                  style: str = "thesis", table_style: str = "three_line") -> Any:
    """Build a table with the specified style.

    table_style options:
    - "three_line": Classic three-line table (三线表), no background
    - "academic": Three-line + shaded header + zebra striping
    - "grid": Full grid with blue header (Office-style)
    """
    if table_style == "academic":
        return _build_academic_table(doc, headers, data_rows, style)
    elif table_style == "grid":
        return _build_grid_table(doc, headers, data_rows, style)
    else:
        return _build_three_line_table(doc, headers, data_rows, style)


def _build_grid_table(doc: Any, headers: list[str], data_rows: list[list[str]],
                     style: str = "thesis") -> Any:
    """Build a full grid table with blue header background (Office-style).

    All borders: 0.5pt. Header: bold, centered, blue background (#4472C4).
    Data: centered, 10.5pt, alternating row shading.
    """
    if Document is None or OxmlElement is None or qn is None or Pt is None:
        num_rows = len(data_rows) + 1
        num_cols = len(headers)
        return doc.add_table(rows=num_rows, cols=num_cols)

    num_rows = len(data_rows) + 1
    num_cols = len(headers)
    table = doc.add_table(rows=num_rows, cols=num_cols)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    # Full grid borders: 0.5pt (sz=4)
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

    # Table width & alignment
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblW)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    existing_jc = tblPr.find(qn("w:jc"))
    if existing_jc is not None:
        tblPr.remove(existing_jc)
    tblPr.append(jc)

    # Header row: bold, white text, blue background
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        for p in cell.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Times New Roman"
        if qn is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        _set_cell_shading(cell, "4472C4")

    # Data rows with zebra striping
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = table.cell(row_idx + 1, col_idx)
            for p in cell.paragraphs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text) if col_idx < len(row_data) else "")
            run.font.size = Pt(10.5)
            run.font.name = "Times New Roman"
            if qn is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if row_idx % 2 == 1:
                _set_cell_shading(cell, "D9E2F3")

    return table


def _build_academic_table(doc: Any, headers: list[str], data_rows: list[list[str]],
                          style: str = "thesis") -> Any:
    """Build a styled academic table with shaded header.

    Three-line borders (top 1.5pt, header separator 0.75pt, bottom 1.5pt)
    plus light gray background on header row for visual clarity.
    Headers: bold, centered, 10.5pt SimSun/Times New Roman, gray background.
    Data: centered, 10.5pt, alternating row shading (zebra stripes).
    """
    if Document is None or OxmlElement is None or qn is None or Pt is None:
        num_rows = len(data_rows) + 1
        num_cols = len(headers)
        return doc.add_table(rows=num_rows, cols=num_cols)

    num_rows = len(data_rows) + 1
    num_cols = len(headers)
    table = doc.add_table(rows=num_rows, cols=num_cols)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    # Remove default borders
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

    # Table width & alignment
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblW)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    existing_jc = tblPr.find(qn("w:jc"))
    if existing_jc is not None:
        tblPr.remove(existing_jc)
    tblPr.append(jc)

    # Three-line borders
    top_border = OxmlElement("w:top")
    top_border.set(qn("w:val"), "single")
    top_border.set(qn("w:sz"), "12")
    top_border.set(qn("w:space"), "0")
    top_border.set(qn("w:color"), "000000")
    borders.append(top_border)

    bottom_border = OxmlElement("w:bottom")
    bottom_border.set(qn("w:val"), "single")
    bottom_border.set(qn("w:sz"), "12")
    bottom_border.set(qn("w:space"), "0")
    bottom_border.set(qn("w:color"), "000000")
    borders.append(bottom_border)

    # Header row: bold, centered, gray background (#D9E2F3 light blue-gray)
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        for p in cell.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.name = "Times New Roman"
        if qn is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        # Header borders + background
        _set_cell_border(cell,
                         top={"val": "single", "sz": "12", "color": "000000"},
                         bottom={"val": "single", "sz": "6", "color": "000000"})
        _set_cell_shading(cell, "D9E2F3")

    # Data rows with zebra striping
    for row_idx, row_data in enumerate(data_rows):
        is_last_row = (row_idx == len(data_rows) - 1)
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = table.cell(row_idx + 1, col_idx)
            for p in cell.paragraphs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text) if col_idx < len(row_data) else "")
            run.font.size = Pt(10.5)
            run.font.name = "Times New Roman"
            if qn is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            # Bottom border on last row
            if is_last_row:
                _set_cell_border(cell, bottom={"val": "single", "sz": "12", "color": "000000"})
            # Zebra striping on even rows
            if row_idx % 2 == 1:
                _set_cell_shading(cell, "F2F2F2")

    return table


def _set_cell_shading(cell: Any, color: str) -> None:
    """Set cell background color."""
    if OxmlElement is None or qn is None:
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shading
    existing_shd = tcPr.find(qn("w:shd"))
    if existing_shd is not None:
        tcPr.remove(existing_shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def _set_cell_border(cell: Any, **kwargs: dict[str, str]) -> None:
    """Set cell border. Usage: _set_cell_border(cell, top={"val": "single", "sz": "4", "color": "000000"})"""
    if OxmlElement is None or qn is None:
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        for attr_name, attr_val in attrs.items():
            element.set(qn(f"w:{attr_name}"), attr_val)


def _apply_inline_formatting(paragraph: Any, text: str, style: str = "thesis") -> None:
    """Parse inline markdown formatting and add runs to paragraph.

    Handles: **bold**, *italic*, `code`, $inline math$, 【Author Year】 citations.
    Scans the entire text left-to-right for the earliest match each iteration.
    """
    if paragraph is None:
        return

    pos = 0
    font_size = Pt(12 if style == "thesis" else 11)
    ea_body = "宋体" if style == "thesis" else "等线"

    while pos < len(text):
        remaining = text[pos:]

        # Find the earliest match among all inline patterns
        bold_m = re.search(r"\*\*(.+?)\*\*", remaining)
        italic_m = re.search(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", remaining)
        code_m = re.search(r"`([^`]+)`", remaining)
        math_m = re.search(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)", remaining)

        matches = []
        if bold_m:
            matches.append(("bold", bold_m.start(), bold_m.end(), bold_m.group(1)))
        if italic_m:
            matches.append(("italic", italic_m.start(), italic_m.end(), italic_m.group(1)))
        if code_m:
            matches.append(("code", code_m.start(), code_m.end(), code_m.group(1)))
        if math_m:
            matches.append(("math", math_m.start(), math_m.end(), math_m.group(1)))

        if not matches:
            if remaining:
                run = paragraph.add_run(remaining)
                run.font.name = "Times New Roman"
                run.font.size = font_size
                if qn is not None:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), ea_body)
            break

        # Pick the earliest match
        matches.sort(key=lambda x: x[1])
        mtype, mstart, mend, mcontent = matches[0]

        # Add plain text before the match
        if mstart > 0:
            run = paragraph.add_run(remaining[:mstart])
            run.font.name = "Times New Roman"
            run.font.size = font_size
            if qn is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), ea_body)

        # Add the formatted run
        if mtype == "bold":
            run = paragraph.add_run(mcontent)
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = font_size
            if qn is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if style == "thesis" else "等线")
        elif mtype == "italic":
            run = paragraph.add_run(mcontent)
            run.font.italic = True
            run.font.name = "Times New Roman"
            run.font.size = font_size
            if qn is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), ea_body)
        elif mtype == "code":
            run = paragraph.add_run(mcontent)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            if qn is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        elif mtype == "math":
            omml = _convert_latex_to_omml(mcontent)
            if omml:
                _insert_omml_formula(paragraph, omml)
            else:
                run = paragraph.add_run(mcontent)
                run.font.name = "Cambria Math"
                run.font.italic = True
                run.font.size = font_size
                if qn is not None:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")

        pos += mend


def _add_text_run(paragraph: Any, text: str, style: str = "thesis") -> None:
    """Add a plain text run with proper font settings."""
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12 if style == "thesis" else 11)
    if qn is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _smart_image_width(image_path: Path, max_width_cm: float = 15.0,
                        max_height_cm: float = 20.0) -> float:
    """Calculate optimal image width in cm based on actual dimensions.

    Ensures the image fits within both max_width_cm and max_height_cm,
    preserving aspect ratio. For thesis style, max_height_cm defaults to 20cm
    (A4 usable area minus margins and caption space).
    """
    try:
        img = PILImage.open(image_path)
        w_px, h_px = img.size
        img.close()
        if h_px == 0 or w_px == 0:
            return max_width_cm
        aspect = w_px / h_px

        # Calculate width that fits height constraint
        width_for_height = max_height_cm * aspect

        # Calculate width based on aspect ratio category
        if aspect >= 2.0:
            # Landscape / panoramic — full width
            preferred_width = max_width_cm
        elif aspect >= 1.0:
            # Slightly landscape — 90% width
            preferred_width = max_width_cm * 0.9
        elif aspect >= 0.6:
            # Slightly portrait — 70% width
            preferred_width = max_width_cm * 0.7
        else:
            # Portrait — 55% width
            preferred_width = max_width_cm * 0.55

        # Use the smaller of preferred width and height-constrained width
        return min(preferred_width, width_for_height)
    except Exception:
        return max_width_cm


def _resolve_image_hint_path(source_path: Path, project_root: Path, hint: str = "") -> Path | None:
    candidate_text = str(hint or "").strip()
    if not candidate_text:
        return None
    candidate = Path(candidate_text)
    probes = [candidate] if candidate.is_absolute() else [
        (source_path.parent / candidate).resolve(),
        (project_root / candidate).resolve(),
    ]
    for probe in probes:
        if probe.exists() and probe.is_file():
            return probe
    return None


def _search_image(source_path: Path, project_root: Path, alt_text: str = "",
                   caption_text: str = "") -> Path | None:
    """Search for an image file across multiple locations.

    Search strategies (in order):
    1. Direct path match (relative to source or common figure dirs)
    2. Filename-based match using figure number patterns
    3. Semantic keyword match using alt_text and caption_text against filenames
    4. Recursive search in all search_dirs (not just project_root)
    """
    # Build candidate names from alt text
    candidates = []
    if alt_text:
        # Extract figure number like "4-1" from "图4-1 ..." or "Figure 4-1 ..."
        fig_match = re.search(r"(\d+[-．.]\d+)", alt_text)
        if fig_match:
            fig_num = fig_match.group(1).replace("．", ".").replace("-", "-")
            candidates.append(f"fig_{fig_num}.png")
            candidates.append(f"figure_{fig_num}.png")
            candidates.append(f"fig{fig_num}.png")
            candidates.append(f"figure{fig_num}.png")
            candidates.append(f"图{fig_num}.png")

    search_dirs = [
        source_path.parent,
        project_root / "output" / "figures",
        Path("G:/sci/output/figures"),
        project_root / "external" / "ros2-slam" / "output" / "figures",
        # Also scan any external/ subdirectories that have output/figures/
        project_root,
    ]
    # Auto-discover additional figure directories under external/
    for ext_dir in (project_root / "external").glob("*/output/figures"):
        if ext_dir.is_dir() and ext_dir not in search_dirs:
            search_dirs.append(ext_dir)

    for search_dir in search_dirs:
        if not search_dir.exists():
            continue
        for candidate in candidates:
            img_path = search_dir / candidate
            if img_path.exists():
                return img_path
        # Also try PNG/JPG variants
        for candidate in candidates:
            base = Path(candidate).stem
            for ext in [".png", ".jpg", ".jpeg", ".bmp", ".svg"]:
                img_path = search_dir / (base + ext)
                if img_path.exists():
                    return img_path

    # Recursive search: filename-based patterns
    if candidates:
        fig_match = re.search(r"(\d+[-．.]\d+)", alt_text) if alt_text else None
        if fig_match:
            fig_num = fig_match.group(1).replace("．", ".")
            specific_patterns = [
                f"fig_{fig_num.replace('-', '_')}",
                f"fig{fig_num.replace('-', '')}",
                f"figure_{fig_num.replace('-', '_')}",
                f"图{fig_num}",
            ]
            for search_dir in search_dirs:
                if not search_dir.exists():
                    continue
                try:
                    for f in search_dir.iterdir():
                        if f.is_file() and f.suffix.lower() in (".png", ".jpg", ".jpeg", ".bmp", ".svg"):
                            stem_lower = f.stem.lower()
                            if any(pat.lower() in stem_lower for pat in specific_patterns):
                                return f
                except (OSError, PermissionError):
                    pass

    # Semantic keyword match: extract keywords from alt_text + caption_text
    # and match against image filenames
    search_text = f"{alt_text} {caption_text}".lower()
    if search_text.strip():
        # Build keyword -> likely filename mappings for common domains
        keyword_map = {
            "slam": ["slam", "建图", "地图", "定位"],
            "navigation": ["navigation", "导航", "planner", "路径", "nav"],
            "comparison": ["comparison", "对比", "比较", "柱状图", "bar"],
            "heatmap": ["heatmap", "热力图"],
            "radar": ["radar", "雷达图"],
            "convergence": ["convergence", "收敛", "曲线"],
            "architecture": ["architecture", "架构", "系统", "结构"],
            "scenario": ["scenario", "场景"],
            "mppi": ["mppi", "控制器", "controller"],
            "sensitivity": ["sensitivity", "敏感性", "参数"],
            "mapping": ["mapping", "建图效果", "地图对比"],
            "trajectory": ["trajectory", "轨迹"],
            "obstacle": ["obstacle", "障碍", "避障", "动态"],
            "resource": ["resource", "资源", "cpu", "内存"],
            "fault": ["fault", "故障", "异常", "日志"],
            "timing": ["timing", "时序"],
            "mode": ["mode", "模式", "切换"],
            "data_flow": ["data_flow", "数据流", "topic"],
            "node": ["node", "节点", "通信"],
            "optimization": ["optimization", "优化", "位姿"],
            "control": ["control", "控制", "闭环"],
            "frontend": ["frontend", "前端"],
            "backend": ["backend", "后端"],
            "recovery": ["recovery", "恢复"],
            "detection": ["detection", "检测"],
            "parameter": ["parameter", "参数", "配置"],
            "startup": ["startup", "启动"],
            "state": ["state", "状态机"],
        }
        matched_keywords = []
        for keyword, triggers in keyword_map.items():
            for trigger in triggers:
                if trigger in search_text:
                    matched_keywords.append(keyword)
                    break

        if matched_keywords:
            best_match = None
            best_score = 0
            for search_dir in search_dirs:
                if not search_dir.exists():
                    continue
                try:
                    for f in search_dir.iterdir():
                        if not f.is_file() or f.suffix.lower() not in (".png", ".jpg", ".jpeg", ".bmp", ".svg"):
                            continue
                        stem_lower = f.stem.lower()
                        score = sum(1 for kw in matched_keywords if kw in stem_lower)
                        if score > best_score:
                            best_score = score
                            best_match = f
                except (OSError, PermissionError):
                    pass

            if best_match and best_score >= 1:
                return best_match

    return None


def _create_reference_docx(output_path: Path) -> bool:
    """Create a pandoc reference DOCX with correct fonts: SimSun (Chinese) + Times New Roman (English)."""
    if Document is None:
        return False
    try:
        doc = Document()

        # Page setup
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        # Set East Asian font to SimSun
        rpr = style.element.get_or_add_rPr()
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:eastAsia'), 'SimSun')
        rfonts.set(qn('w:ascii'), 'Times New Roman')
        rfonts.set(qn('w:hAnsi'), 'Times New Roman')
        rpr.insert(0, rfonts)

        # Configure heading styles
        for level in range(1, 5):
            try:
                heading_style = doc.styles[f'Heading {level}']
                hfont = heading_style.font
                hfont.name = 'Times New Roman'
                hfont.color.rgb = RGBColor(0, 0, 0)
                hfont.bold = True
                sizes = {1: Pt(22), 2: Pt(16), 3: Pt(14), 4: Pt(12)}
                if level in sizes:
                    hfont.size = sizes[level]
                hrpr = heading_style.element.get_or_add_rPr()
                hrfonts = OxmlElement('w:rFonts')
                hrfonts.set(qn('w:eastAsia'), 'SimHei')
                hrfonts.set(qn('w:ascii'), 'Times New Roman')
                hrfonts.set(qn('w:hAnsi'), 'Times New Roman')
                hrpr.insert(0, hrfonts)
            except KeyError:
                pass

        # Write placeholder content so pandoc can pick up styles
        doc.add_heading('Heading 1', level=1)
        doc.add_paragraph('Body text placeholder.')
        doc.add_heading('Heading 2', level=2)
        doc.add_paragraph('Body text placeholder.')
        doc.add_heading('Heading 3', level=3)
        doc.add_paragraph('Body text placeholder.')

        doc.save(str(output_path))
        return True
    except Exception:
        return False


_REFERENCE_DOCX_PATH = Path(__file__).parent / "_pandoc_reference.docx"


def _ensure_reference_docx() -> Path | None:
    """Ensure the reference DOCX exists, creating it if needed."""
    if _REFERENCE_DOCX_PATH.exists():
        return _REFERENCE_DOCX_PATH
    if _create_reference_docx(_REFERENCE_DOCX_PATH):
        return _REFERENCE_DOCX_PATH
    return None


def _convert_docx_tables_to_three_line(docx_path: Path) -> None:
    """Post-process a DOCX to convert all tables to three-line (三线表) format."""
    if Document is None or OxmlElement is None or qn is None:
        return
    try:
        doc = Document(str(docx_path))
        for table in doc.tables:
            _apply_three_line_borders(table)
        doc.save(str(docx_path))
    except Exception:
        pass


def _apply_three_line_borders(table: Any) -> None:
    """Apply three-line table borders to an existing python-docx table."""
    if OxmlElement is None or qn is None:
        return
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    # Remove existing borders
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    borders = OxmlElement("w:tblBorders")

    # Top border: 1.5pt (sz=12)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), "000000")
    borders.append(top)

    # Bottom border: 1.5pt (sz=12)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)

    # No left/right/insideV borders
    for name in ("left", "right", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)

    # No insideH (no intermediate horizontal lines)
    insideH = OxmlElement("w:insideH")
    insideH.set(qn("w:val"), "none")
    insideH.set(qn("w:sz"), "0")
    insideH.set(qn("w:space"), "0")
    insideH.set(qn("w:color"), "auto")
    borders.append(insideH)

    tblPr.append(borders)

    # Add bottom border to first row (header separator): 0.75pt (sz=6)
    if table.rows:
        first_row = table.rows[0]
        for cell in first_row.cells:
            _set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": "000000"})

    # Format header cells: bold, centered, 10.5pt
    if table.rows and Pt is not None:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if WD_ALIGN_PARAGRAPH else None
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10.5)
                    run.font.name = "Times New Roman"
                    if qn is not None:
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _try_pandoc_export(source_path: Path, target_path: Path, style: str, project_root: Path | None = None) -> dict[str, Any] | None:
    """Attempt to export via pandoc for proper LaTeX math rendering in DOCX."""
    import shutil
    import subprocess

    pandoc = shutil.which("pandoc")
    if not pandoc:
        return None

    # Pre-process: remove \tag{} from LaTeX (pandoc doesn't support it)
    content = source_path.read_text(encoding="utf-8")

    cleaned = re.sub(r" \\tag\{[^}]+\}", "", content)

    # Resolve image paths: convert relative paths to absolute and prefer PNG over PDF
    # pandoc cannot embed PDF images, so we replace .pdf with .png
    source_dir = source_path.resolve().parent
    # Also try project root (for cases where source is in drafts/refinement/)
    project_dir = Path(project_root).resolve() if project_root else source_dir

    def _resolve_image_path(m: re.Match) -> str:
        caption = m.group(1)
        path_str = m.group(2)
        # Try source_dir first, then project_dir
        for base in [source_dir, project_dir]:
            img_path = base / path_str
            if img_path.exists():
                png_path = img_path.with_suffix(".png")
                if png_path.exists():
                    resolved = png_path
                elif img_path.suffix.lower() == ".pdf":
                    continue  # skip PDF, try next base
                else:
                    resolved = img_path
                return f"![{caption}]({resolved.as_posix()})"
        return m.group(0)  # keep original if not found

    cleaned = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", _resolve_image_path, cleaned)

    # Write temp cleaned file
    tmp = source_path.with_suffix(".pandoc-tmp.md")
    try:
        tmp.write_text(cleaned, encoding="utf-8")

        cmd = [
            pandoc, str(tmp),
            "-o", str(target_path),
            "--from", "markdown+tex_math_dollars",
            "--to", "docx",
            "--standalone",
        ]
        # Use reference DOCX for correct fonts (SimSun + Times New Roman)
        ref_docx = _ensure_reference_docx()
        if ref_docx:
            cmd.extend(["--reference-doc", str(ref_docx)])
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            return None

        # Post-process: force black font on all headings via python-docx
        if Document is not None and target_path.exists():
            _force_docx_heading_colors(target_path)
            # Convert tables to three-line format
            _convert_docx_tables_to_three_line(target_path)

        return {
            "kind": "docx_export",
            "artifact": "paper",
            "source_path": str(source_path),
            "output_path": str(target_path),
            "exported_at": datetime.now(timezone.utc).isoformat(),
            "style": style,
            "engine": "pandoc",
        }
    except Exception:
        return None
    finally:
        tmp.unlink(missing_ok=True)


def _force_docx_heading_colors(docx_path: Path) -> None:
    """Ensure all headings in a DOCX have black font color."""
    if Document is None:
        return
    doc = Document(str(docx_path))
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name and paragraph.style.name.startswith("Heading"):
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
    doc.save(str(docx_path))


def _validate_image_placement(
    blocks: list[dict[str, Any]],
    archetype: str = "engineering",
) -> list[dict[str, str]]:
    """Validate image placement at DOCX export time.

    Checks:
    1. Theory chapter images are principle/design, not result/comparison
    2. Experiment chapter images follow scene->process->result->comparison order
    3. Conclusion chapter has no images (or only summary)

    Returns list of warnings: [{"chapter": int, "figure": str, "issue": str}]
    """
    try:
        from image_roles import classify_image, is_role_compatible, get_image_order_key, CHAPTER_ROLE_MATRIX
    except ImportError:
        return []

    warnings: list[dict[str, str]] = []
    current_chapter = 0
    chapter_images: dict[int, list[dict[str, Any]]] = {}

    for block in blocks:
        if block.get("type") == "heading":
            text = block.get("text", "")
            ch_match = re.match(r"(\d+)\.", text)
            if ch_match:
                current_chapter = int(ch_match.group(1))

        if block.get("type") in ("image", "image_placeholder"):
            ref = block.get("alt", block.get("ref", ""))
            path_str = block.get("path", "")
            name = Path(path_str).stem if path_str else ref
            role = classify_image(name, alt_text=ref)
            chapter_images.setdefault(current_chapter, []).append({
                "role": role,
                "ref": ref or name,
                "order": get_image_order_key(role),
            })

    matrix = CHAPTER_ROLE_MATRIX.get(archetype, CHAPTER_ROLE_MATRIX.get("engineering", {}))

    for ch, images in chapter_images.items():
        ch_info = matrix.get(ch, {})
        allowed = ch_info.get("allowed", set())

        for img in images:
            if allowed and img["role"] not in allowed:
                warnings.append({
                    "chapter": str(ch),
                    "figure": img["ref"],
                    "issue": f"Role '{img['role']}' not ideal for chapter {ch} (allowed: {', '.join(sorted(allowed))})",
                })

        # Check ordering within chapter
        if len(images) > 1 and ch_info.get("order"):
            orders = [img["order"] for img in images]
            if orders != sorted(orders):
                warnings.append({
                    "chapter": str(ch),
                    "figure": "ordering",
                    "issue": "Images not in prescribed order (scene->process->result->comparison)",
                })

    return warnings


def export_markdown_to_docx(
    project_root: str | Path = ".",
    artifact: str = "paper",
    source: str | Path | None = None,
    output_path: str | Path | None = None,
    style: str = "thesis",
    table_style: str = "three_line",
    archetype: str = "engineering",
) -> dict[str, Any]:
    root = Path(project_root).resolve()
    source_path = _resolve_source(root, artifact, source)
    target_path = Path(output_path).resolve() if output_path else _default_output_path(root, source_path, ".docx")

    # Use python-docx as primary path (handles compressed tables, formula conversion,
    # three-line tables, inline formatting, image placeholders, reference formatting).
    # Pandoc is available as fallback for edge cases.
    if Document is None:
        # Fall back to pandoc if python-docx is unavailable
        pandoc_result = _try_pandoc_export(source_path, target_path, style, project_root=root)
        if pandoc_result is not None:
            try:
                word_cross_ref_process(target_path, target_path)
                pandoc_result["cross_ref_applied"] = True
            except Exception:
                pandoc_result["cross_ref_applied"] = False
            return pandoc_result
        raise RuntimeError("Neither python-docx nor pandoc is available; DOCX export is unavailable.")

    lines = _markdown_lines(source_path)
    blocks = _iter_markdown_blocks(lines)

    # Validate image placement (advisory warnings, not blocking)
    placement_warnings = _validate_image_placement(blocks, archetype=archetype)
    if placement_warnings:
        print(f"[Export] Image placement warnings ({len(placement_warnings)}):")
        for w in placement_warnings:
            print(f"  Ch{w['chapter']}: {w['figure']} - {w['issue']}")

    # Run full paper quality validation
    quality_report = None
    try:
        from paper_quality import validate_paper, format_report
        quality_report = validate_paper(source_path, language="zh")
        if quality_report["issues"]:
            print(f"[Export] Paper quality score: {quality_report['score']}/100")
            high_issues = [i for i in quality_report["issues"] if i["severity"] == "high"]
            if high_issues:
                print(f"[Export WARNING] 论文存在{len(high_issues)}个关键缺陷，建议修复后再导出：")
                for issue in high_issues:
                    print(f"  [HIGH] {issue['category']}: {issue['message']}")
            else:
                for issue in quality_report["issues"]:
                    if issue["severity"] == "medium":
                        print(f"  [MEDIUM] {issue['category']}: {issue['message']}")
    except ImportError:
        pass

    document = Document()
    _apply_docx_style(document, style)
    title_written = False

    # Detect the minimum heading level in the document to map correctly.
    # If the doc starts with ## (no H1), then ## = chapter (Heading 1),
    # ### = section (Heading 2), etc.
    min_heading_level = min(
        (b["level"] for b in blocks if b.get("type") == "heading"),
        default=1,
    )
    heading_offset = min_heading_level - 1  # e.g. min=2 → offset=1

    # Track figure/table numbers for caption matching
    last_table_num = 0
    last_figure_num = 0
    used_image_paths: set[str] = set()  # Track already-assigned images to avoid duplicates
    used_image_count: dict[str, int] = {}  # Track usage count per image (allow cross-chapter reuse)

    # Pre-load all available images for fallback sequential assignment
    all_available_images: list[Path] = []
    img_search_dirs = [
        root / "output" / "figures",
        Path("G:/sci/output/figures"),
    ]
    for ext_dir in (root / "external").glob("*/output/figures"):
        if ext_dir.is_dir() and ext_dir not in img_search_dirs:
            img_search_dirs.append(ext_dir)
    seen_img_paths: set[str] = set()
    for sd in img_search_dirs:
        if not sd.exists():
            continue
        try:
            for f in sd.iterdir():
                if f.is_file() and f.suffix.lower() in (".png", ".jpg", ".jpeg", ".bmp", ".svg"):
                    resolved = str(f.resolve())
                    if resolved not in seen_img_paths:
                        seen_img_paths.add(resolved)
                        all_available_images.append(f)
        except (OSError, PermissionError):
            pass

    _skip_caption = False  # flag to skip caption after filtered AI image

    for block_idx, block in enumerate(blocks):
        btype = block.get("type", "")

        if btype == "heading":
            level = min(int(block["level"]), 9)
            text = block["text"]
            # Clean formula syntax descriptions from headings (e.g. "$$...$$")
            text = re.sub(r'\$\$[^$]+\$\$', '', text).strip()
            # Map to Word heading level: offset so the top-level
            # markdown heading maps to Heading 1 (chapter).
            doc_level = max(level - heading_offset, 1)
            if doc_level > 4:
                doc_level = 4  # cap at Heading 4

            if title_written and _should_insert_page_break(doc_level, text):
                document.add_page_break()
            # Only use Title style for original H1 (# Title),
            # not for mapped chapter headings (## 1. 绪论 → doc_level 1)
            if not title_written and level == 1 and heading_offset == 0:
                # Strip draft markers from title (e.g. "论文写作稿：" prefix)
                text = re.sub(r'^(?:论文写作稿|Paper Draft)\s*[：:]\s*', '', text)
                paragraph = document.add_heading(text, level=0)
                _format_heading(paragraph, 0, style)
                title_written = True
            else:
                paragraph = document.add_heading(text, level=doc_level)
                _format_heading(paragraph, doc_level, style)

        elif btype == "paragraph":
            paragraph = document.add_paragraph()
            _format_body_paragraph(
                paragraph,
                style,
                first_line_indent=not _is_front_matter_line(block["text"]),
            )
            _apply_inline_formatting(paragraph, block["text"], style)

        elif btype == "formula_paragraph":
            # Paragraph containing display formulas — extract and render them.
            # Strategy: consecutive text+inline_formula segments merge into one paragraph;
            # display_formula segments each get their own centered paragraph.
            segments = _extract_formulas_from_text(block["text"])
            i = 0
            while i < len(segments):
                seg = segments[i]
                if seg["type"] == "display_formula":
                    # Display formula: centered, with optional right-aligned tag
                    # 参考论文格式: 公式居中, 编号"（式X.Y）"在右侧
                    omml = _convert_latex_to_omml(seg["content"])
                    if omml:
                        formula_para = document.add_paragraph()
                        formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        formula_para.paragraph_format.space_before = Pt(3)
                        formula_para.paragraph_format.space_after = Pt(3)
                        formula_para.paragraph_format.first_line_indent = Cm(0) if Cm else 0
                        formula_para.paragraph_format.line_spacing = 1.5
                        _insert_omml_formula(formula_para, omml)
                        # Add tag number right-aligned: "（式X.Y）"
                        tag_text = _display_equation_tag_text(seg.get("tag"))
                        if tag_text:
                            tag_run = formula_para.add_run(f"\t{tag_text}")
                            tag_run.font.name = "Times New Roman"
                            tag_run.font.size = Pt(12 if style == "thesis" else 11)
                            if qn is not None:
                                tag_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    else:
                        # Fallback: render as styled text
                        formula_text = seg["content"]
                        tag_text = _display_equation_tag_text(seg.get("tag"))
                        if tag_text:
                            formula_text = f"{formula_text}\t{tag_text}"
                        formula_para = document.add_paragraph(formula_text)
                        formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        formula_para.paragraph_format.space_before = Pt(3)
                        formula_para.paragraph_format.space_after = Pt(3)
                        formula_para.paragraph_format.first_line_indent = Cm(0) if Cm else 0
                        formula_para.paragraph_format.line_spacing = 1.5
                        for run in formula_para.runs:
                            run.font.name = "Cambria Math"
                            run.font.italic = True
                            run.font.size = Pt(12 if style == "thesis" else 11)
                            if qn is not None:
                                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
                    i += 1
                else:
                    # Collect consecutive text + inline_formula segments into one paragraph
                    para = document.add_paragraph()
                    # Determine first_line_indent from the first text segment
                    first_text_seg = seg
                    is_front = _is_front_matter_line(seg.get("content", ""))
                    _format_body_paragraph(para, style, first_line_indent=not is_front)
                    while i < len(segments) and segments[i]["type"] in ("text", "inline_formula"):
                        s = segments[i]
                        if s["type"] == "text":
                            _apply_inline_formatting(para, s["content"], style)
                        elif s["type"] == "inline_formula":
                            omml = _convert_latex_to_omml(s["content"])
                            if omml:
                                _insert_omml_formula(para, omml)
                            else:
                                run = para.add_run(s["content"])
                                run.font.name = "Cambria Math"
                                run.font.italic = True
                                run.font.size = Pt(12 if style == "thesis" else 11)
                                if qn is not None:
                                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
                        i += 1

        elif btype == "table":
            rows = block["rows"]
            body_rows = rows
            if len(rows) >= 2 and _table_divider("| " + " | ".join(rows[1]) + " |"):
                body_rows = [rows[0], *rows[2:]]
            if not body_rows:
                continue
            headers = body_rows[0]
            data = body_rows[1:]
            # Normalize column count
            max_cols = max(len(r) for r in body_rows)
            headers = headers + [""] * (max_cols - len(headers))
            data = [r + [""] * (max_cols - len(r)) for r in data]
            _build_table(document, headers, data, style, table_style)
            # Track table number
            last_table_num += 1

        elif btype == "table_caption":
            caption_para = document.add_paragraph(block["text"])
            _format_caption_paragraph(caption_para, style)
            for run in caption_para.runs:
                run.font.size = Pt(10.5)
                run.font.name = "Times New Roman"
                if qn is not None:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            last_table_num += 1

        elif btype == "figure_caption":
            if _skip_caption:
                _skip_caption = False
                continue
            caption_para = document.add_paragraph(block["text"])
            _format_caption_paragraph(caption_para, style)
            for run in caption_para.runs:
                run.font.size = Pt(10.5)
                run.font.name = "Times New Roman"
                if qn is not None:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            last_figure_num += 1

        elif btype == "image_placeholder":
            # [此处插入图X-Y] placeholder — try to find actual image
            ref = str(block.get("ref", "") or "").strip()
            placeholder_caption = str(block.get("caption", "") or "").strip()
            placeholder_type = str(block.get("figure_type", "") or "").strip()
            placeholder_goal = str(block.get("goal", "") or "").strip()
            placeholder_evidence = str(block.get("evidence", "") or "").strip()
            # Look ahead for figure caption text to improve image search.
            # Search both dedicated figure_caption blocks AND paragraphs that
            # contain the figure number (captions often mixed with explanation text)
            caption_text = placeholder_caption
            for future_block in blocks[block_idx + 1 : block_idx + 8]:
                ftype = future_block.get("type", "")
                ftext = future_block.get("text", "")
                if ftype == "figure_caption":
                    caption_text = ftext
                    break
                # Also check paragraph blocks for figure number pattern
                if ref and ftype in ("paragraph", "formula_paragraph") and f"图{ref}" in ftext:
                    # Extract just the caption part (first sentence or up to 100 chars)
                    caption_match = re.search(rf"图{ref}\s*[^\n。]*", ftext)
                    if caption_match:
                        caption_text = caption_match.group(0)
                    else:
                        caption_text = ftext[:100]
                    break
            search_term = f"图{ref}" if ref else placeholder_caption
            img_path = _resolve_image_hint_path(source_path, root, placeholder_evidence)
            if img_path is None:
                img_path = _search_image(
                    source_path,
                    root,
                    search_term,
                    caption_text=" ".join(part for part in [caption_text, placeholder_goal, placeholder_type, placeholder_evidence] if part).strip(),
                )
            # Allow cross-chapter reuse but cap at 2 uses per image
            if img_path:
                resolved = str(img_path.resolve())
                count = used_image_count.get(resolved, 0)
                if count >= 2:
                    img_path = None  # Overused, try fallback
            if img_path and img_path.exists():
                used_image_paths.add(str(img_path.resolve()))
                used_image_count[str(img_path.resolve())] = used_image_count.get(str(img_path.resolve()), 0) + 1
                try:
                    img_w = _smart_image_width(img_path, max_width_cm=(13.5 if style == "thesis" else 15.0), max_height_cm=(18.0 if style == "thesis" else 20.0))
                    document.add_picture(str(img_path), width=Cm(img_w))
                    pic_para = document.paragraphs[-1]
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    para = document.add_paragraph(f"[图片加载失败: {img_path.name}]")
                # Add caption
                clean_caption = re.sub(r"^(?:图|Figure)\s*\d+(?:[-．.]\d+)?\s*", "", caption_text).strip()
                cap_text = f"图{ref}" if ref else (clean_caption or placeholder_caption or "图片")
                if ref and clean_caption:
                    cap_text = f"{cap_text} {clean_caption}"
                cap_para = document.add_paragraph(cap_text)
                _format_caption_paragraph(cap_para, style)
                for run in cap_para.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = "Times New Roman"
                    if qn is not None:
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                last_figure_num += 1
            else:
                # Fallback: sequentially assign an unused image
                fallback_img = None
                for avail in all_available_images:
                    resolved = str(avail.resolve())
                    count = used_image_count.get(resolved, 0)
                    if count < 1 and resolved not in used_image_paths:
                        fallback_img = avail
                        break
                if fallback_img and fallback_img.exists():
                    used_image_paths.add(str(fallback_img.resolve()))
                    used_image_count[str(fallback_img.resolve())] = 1
                    try:
                        img_w = _smart_image_width(fallback_img, max_width_cm=(13.5 if style == "thesis" else 15.0))
                        document.add_picture(str(fallback_img), width=Cm(img_w))
                        pic_para = document.paragraphs[-1]
                        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        para = document.add_paragraph(f"[图片加载失败: {fallback_img.name}]")
                    clean_caption = re.sub(r"^(?:图|Figure)\s*\d+(?:[-．.]\d+)?\s*", "", caption_text).strip()
                    cap_text = f"图{ref}" if ref else (clean_caption or placeholder_caption or "图片")
                    if ref and clean_caption:
                        cap_text = f"{cap_text} {clean_caption}"
                    cap_para = document.add_paragraph(cap_text)
                    _format_caption_paragraph(cap_para, style)
                    for run in cap_para.runs:
                        run.font.size = Pt(10.5)
                        run.font.name = "Times New Roman"
                        if qn is not None:
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    last_figure_num += 1
                else:
                    # Preserve a visible placeholder with enough guidance for manual补图.
                    placeholder_lines = [f"[此处插入图{ref}]" if ref else f"[待补图] {placeholder_caption or '请补充与正文对应的图片'}"]
                    if placeholder_caption and ref:
                        placeholder_lines.append(f"图{ref} {placeholder_caption}")
                    if placeholder_type:
                        placeholder_lines.append(f"图型建议：{placeholder_type}")
                    if placeholder_goal:
                        placeholder_lines.append(f"应展示内容：{placeholder_goal}")
                    if placeholder_evidence:
                        placeholder_lines.append(f"推荐素材来源：{placeholder_evidence}")
                    for line_text in placeholder_lines:
                        para = document.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run(line_text)
                        run.font.size = Pt(10.5)
                        run.font.color.rgb = RGBColor(169, 169, 169)
                        run.font.name = "Times New Roman"
                        if qn is not None:
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        elif btype == "image":
            image_path = Path(block["path"])
            # Skip AI-generated placeholder images that don't come from the project itself.
            # Real project screenshots have specific naming (household_demo_*, roomsketcher_*, etc.)
            # while AI-generated ones use generic names like system-architecture, data-flow, etc.
            stem = image_path.stem
            _ai_placeholder_patterns = (
                "architecture-diagram", "flowchart-diagram",
            )
            if stem.startswith("gen_ai_"):
                _skip_caption = True
                continue
            if not image_path.is_absolute():
                # Try relative to source dir first
                image_path = (source_path.parent / image_path).resolve()
            if not image_path.exists():
                # Try relative to project root
                image_path = (root / block["path"]).resolve()
            if not image_path.exists():
                # The alt text may have been translated to Chinese by _cleanup_markdown,
                # but the actual file on disk still uses the original English name.
                # Try to find the original file by searching figure dirs.
                found = _search_image(source_path, root, block.get("alt", ""),
                                       caption_text=block.get("alt", ""))
                if found and str(found.resolve()) not in used_image_paths:
                    image_path = found
                elif found:
                    image_path = Path("")  # Already used, skip
                else:
                    # Last resort: scan project_root/output/figures/ for any PNG
                    figures_dir = root / "output" / "figures"
                    if figures_dir.exists():
                        # Try to find by looking for files that contain the alt text
                        # or use sequential assignment
                        for avail in sorted(figures_dir.glob("*.png")):
                            if str(avail.resolve()) not in used_image_paths:
                                image_path = avail
                                break
                        else:
                            image_path = Path("")
                    else:
                        image_path = Path("")
            if not image_path.exists():
                para = document.add_paragraph(f"[Missing image: {block['path']}]")
                _format_body_paragraph(para, style, first_line_indent=False)
                continue
            used_image_paths.add(str(image_path.resolve()))
            try:
                img_w = _smart_image_width(image_path, max_width_cm=(13.5 if style == "thesis" else 15.0), max_height_cm=(18.0 if style == "thesis" else 20.0))
                # Use document.add_picture() — run.add_picture() has known bugs
                # (python-docx issues #981, #1063, #824)
                document.add_picture(str(image_path), width=Cm(img_w))
                # Format the paragraph that add_picture just created
                pic_para = document.paragraphs[-1]
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                para = document.add_paragraph(f"[图片加载失败: {image_path.name}]")
            if block.get("alt"):
                caption = document.add_paragraph(block["alt"])
                _format_caption_paragraph(caption, style)
                for run in caption.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = "Times New Roman"
                    if qn is not None:
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        elif btype == "reference":
            # Clean up reference text: convert "1. Author..." to "[1] Author..."
            ref_text = block["text"]
            ref_text = re.sub(r"^(\d+)\s*[.．]\s*", r"[\1] ", ref_text)
            para = document.add_paragraph(ref_text)
            # 参考论文: 宋体 12pt 两端对齐 首行缩进2字符
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Cm(0.74) if Cm else 0
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
                if qn is not None:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        elif btype == "bullet":
            if "$$" in block["text"]:
                segments = _extract_formulas_from_text(block["text"])
                i = 0
                while i < len(segments):
                    seg = segments[i]
                    if seg["type"] == "display_formula":
                        omml = _convert_latex_to_omml(seg["content"])
                        if omml:
                            fp = document.add_paragraph()
                            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            fp.paragraph_format.space_before = Pt(6)
                            fp.paragraph_format.space_after = Pt(6)
                            fp.paragraph_format.first_line_indent = Cm(0) if Cm else 0
                            _insert_omml_formula(fp, omml)
                            tag_text = _display_equation_tag_text(seg.get("tag"))
                            if tag_text:
                                tr = fp.add_run(f"\t{tag_text}")
                                tr.font.name = "Times New Roman"
                                tr.font.size = Pt(12 if style == "thesis" else 11)
                                if qn is not None:
                                    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                        else:
                            ft = seg["content"]
                            tag_text = _display_equation_tag_text(seg.get("tag"))
                            if tag_text:
                                ft = f"{ft}\t{tag_text}"
                            fp = document.add_paragraph(ft)
                            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            fp.paragraph_format.first_line_indent = Cm(0) if Cm else 0
                            for r in fp.runs:
                                r.font.name = "Cambria Math"
                                r.font.italic = True
                        i += 1
                    else:
                        para = document.add_paragraph(style="List Bullet")
                        _format_body_paragraph(para, style, first_line_indent=False)
                        while i < len(segments) and segments[i]["type"] in ("text", "inline_formula"):
                            s = segments[i]
                            if s["type"] == "text":
                                _apply_inline_formatting(para, s["content"], style)
                            elif s["type"] == "inline_formula":
                                omml = _convert_latex_to_omml(s["content"])
                                if omml:
                                    _insert_omml_formula(para, omml)
                                else:
                                    r = para.add_run(s["content"])
                                    r.font.name = "Cambria Math"
                                    r.font.italic = True
                            i += 1
            else:
                paragraph = document.add_paragraph(block["text"], style="List Bullet")
                _format_body_paragraph(paragraph, style, first_line_indent=False)

        elif btype == "numbered":
            if "$$" in block["text"]:
                segments = _extract_formulas_from_text(block["text"])
                i = 0
                while i < len(segments):
                    seg = segments[i]
                    if seg["type"] == "display_formula":
                        omml = _convert_latex_to_omml(seg["content"])
                        if omml:
                            fp = document.add_paragraph()
                            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            fp.paragraph_format.space_before = Pt(6)
                            fp.paragraph_format.space_after = Pt(6)
                            fp.paragraph_format.first_line_indent = Cm(0) if Cm else 0
                            _insert_omml_formula(fp, omml)
                            tag_text = _display_equation_tag_text(seg.get("tag"))
                            if tag_text:
                                tr = fp.add_run(f"\t{tag_text}")
                                tr.font.name = "Times New Roman"
                                tr.font.size = Pt(12 if style == "thesis" else 11)
                                if qn is not None:
                                    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                        else:
                            ft = seg["content"]
                            tag_text = _display_equation_tag_text(seg.get("tag"))
                            if tag_text:
                                ft = f"{ft}\t{tag_text}"
                            fp = document.add_paragraph(ft)
                            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            fp.paragraph_format.first_line_indent = Cm(0) if Cm else 0
                            for r in fp.runs:
                                r.font.name = "Cambria Math"
                                r.font.italic = True
                        i += 1
                    else:
                        para = document.add_paragraph(style="List Number")
                        _format_body_paragraph(para, style, first_line_indent=False)
                        while i < len(segments) and segments[i]["type"] in ("text", "inline_formula"):
                            s = segments[i]
                            if s["type"] == "text":
                                _apply_inline_formatting(para, s["content"], style)
                            elif s["type"] == "inline_formula":
                                omml = _convert_latex_to_omml(s["content"])
                                if omml:
                                    _insert_omml_formula(para, omml)
                                else:
                                    r = para.add_run(s["content"])
                                    r.font.name = "Cambria Math"
                                    r.font.italic = True
                            i += 1
            else:
                paragraph = document.add_paragraph(block["text"], style="List Number")
                _format_body_paragraph(paragraph, style, first_line_indent=False)

        elif btype == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            if qn is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")

        elif btype == "pagebreak":
            document.add_page_break()

        elif btype == "toc":
            _insert_toc(document, style)

    # Only add filename-based title if the file has no H1 heading at all
    # and the filename looks like a real title (not "paper-draft")
    if not title_written:
        stem = source_path.stem
        # Skip generic/template filenames
        skip_names = {"paper-draft", "proposal-draft", "literature-review",
                      "research-answer", "research-presentation", "test_generic"}
        if stem.lower() not in skip_names:
            paragraph = document.add_heading(stem, level=0)
            _format_heading(paragraph, 0, style)

    # Post-processing: remove empty paragraphs that produce blank lines in DOCX
    _remove_empty_paragraphs(document)

    target_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(target_path)

    payload = {
        "kind": "docx_export",
        "artifact": artifact,
        "source_path": str(source_path),
        "output_path": str(target_path),
        "exported_at": _now_iso(),
        "block_count": len(blocks),
        "style": style,
    }
    json_path = target_path.with_suffix(".docx.json")
    _write_json(json_path, payload)
    payload["json_path"] = str(json_path)
    # Auto-run cross-reference processing
    try:
        word_cross_ref_process(target_path, target_path)
        payload["cross_ref_applied"] = True
    except Exception:
        payload["cross_ref_applied"] = False
    return payload


def _apply_docx_style(document: Any, style: str) -> None:
    """Apply unified Chinese academic thesis formatting.

    Reference standard: GB/T 7713 + common university thesis templates.

    Font/size mapping (thesis style):
    - Title: 黑体/TNR, 小二(18pt), centered
    - H1 (章): 黑体/TNR, 三号(16pt), centered, 段前12pt 段后6pt
    - H2 (节): 黑体/TNR, 四号(14pt), left-aligned, 段前6pt 段后3pt
    - H3: 黑体/TNR, 小四(12pt), left-aligned, 段前3pt
    - Body: 宋体/TNR, 小四(12pt), justified, first-line indent 2 chars (0.74cm), line spacing 22pt fixed
    - Caption: 宋体/TNR, 五号(10.5pt), centered
    - References: 宋体/TNR, 五号(10.5pt), hanging indent, 1.25x line spacing
    - Page margins: top 2.5cm, bottom 2.5cm, left 3.0cm, right 2.5cm
    """
    if Document is None or Pt is None or Cm is None or qn is None or WD_SECTION is None:
        return

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    if style == "thesis":
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
    else:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Normal style (body text) — 对齐参考论文: 宋体+TNR 12pt 两端对齐
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    body_size = 12 if style == "thesis" else 11
    normal_style.font.size = Pt(body_size)
    # 参考论文: eastAsia=宋体 (不是微软雅黑)
    ea_font = "宋体" if style == "thesis" else "等线"
    rfonts = normal_style._element.rPr.rFonts
    rfonts.set(qn("w:eastAsia"), ea_font)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if style == "thesis":
        # 参考论文: line=360 lineRule=auto (1.5倍行距)
        normal_style.paragraph_format.line_spacing = 1.5
        # 参考论文: firstLineChars=200 (2字符缩进)
        normal_style.paragraph_format.first_line_indent = Cm(0.74)

    # Heading 1 (章标题): 参考论文"正文 一级标题" — 黑体 18pt (sz=36半磅) 两端对齐
    h1 = document.styles["Heading 1"]
    h1.font.name = "黑体"
    h1.font.bold = True
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0, 0, 0)
    rfonts_h1 = h1._element.rPr.rFonts
    rfonts_h1.set(qn("w:eastAsia"), "黑体" if style == "thesis" else "等线")
    rfonts_h1.set(qn("w:ascii"), "黑体")
    rfonts_h1.set(qn("w:hAnsi"), "黑体")
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.line_spacing = 1.5

    # Heading 2 (节标题): 参考论文"正文 二级标题" — 黑体 15pt (sz=30) 两端对齐
    h2 = document.styles["Heading 2"]
    h2.font.name = "黑体"
    h2.font.bold = False  # 参考论文标题不加粗
    h2.font.size = Pt(15)
    h2.font.color.rgb = RGBColor(0, 0, 0)
    rfonts_h2 = h2._element.rPr.rFonts
    rfonts_h2.set(qn("w:eastAsia"), "黑体" if style == "thesis" else "等线")
    rfonts_h2.set(qn("w:ascii"), "黑体")
    rfonts_h2.set(qn("w:hAnsi"), "黑体")
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.line_spacing = 1.5

    # Heading 3 (小节标题): 参考论文"正文 三级标题" — 黑体 14pt (sz=28) 两端对齐
    h3 = document.styles["Heading 3"]
    h3.font.name = "黑体"
    h3.font.bold = False  # 参考论文标题不加粗
    h3.font.size = Pt(14)
    h3.font.color.rgb = RGBColor(0, 0, 0)
    rfonts_h3 = h3._element.rPr.rFonts
    rfonts_h3.set(qn("w:eastAsia"), "黑体" if style == "thesis" else "等线")
    rfonts_h3.set(qn("w:ascii"), "黑体")
    rfonts_h3.set(qn("w:hAnsi"), "黑体")
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h3.paragraph_format.space_before = Pt(3)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.line_spacing = 1.5

    # Heading 4 (if used): 小四(12pt)
    try:
        h4 = document.styles["Heading 4"]
        h4.font.name = "Times New Roman"
        h4.font.bold = True
        h4.font.size = Pt(12)
        h4.font.color.rgb = RGBColor(0, 0, 0)
        h4._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if style == "thesis" else "等线")
    except KeyError:
        pass

    # Title style (level=0 heading): 黑体, 小二(18pt), centered
    try:
        title_style = document.styles["Title"]
        title_style.font.name = "Times New Roman"
        title_style.font.bold = True
        title_style.font.size = Pt(18)
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if style == "thesis" else "等线")
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_before = Pt(24)
        title_style.paragraph_format.space_after = Pt(18)
    except KeyError:
        pass

    _ensure_footer_page_number(section)


def _format_heading(paragraph: Any, level: int, style: str) -> None:
    if WD_ALIGN_PARAGRAPH is None:
        return
    if level == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(18)
    elif level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(6 if level == 2 else 3)
        paragraph.paragraph_format.space_after = Pt(3 if level == 2 else 0)

    # Explicitly set font on runs (add_heading doesn't inherit style settings)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        if qn is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if style == "thesis" else "等线")
        if Pt is not None:
            sizes = {0: 18, 1: 16, 2: 14, 3: 12}
            if level in sizes:
                run.font.size = Pt(sizes[level])


def _format_body_paragraph(paragraph: Any, style: str, first_line_indent: bool = True) -> None:
    if WD_ALIGN_PARAGRAPH is None or Pt is None or Cm is None:
        return
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if first_line_indent else Cm(0)


def _format_caption_paragraph(paragraph: Any, style: str) -> None:
    if WD_ALIGN_PARAGRAPH is None or Pt is None or Cm is None:
        return
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Cm(0)


def _insert_toc(document: Any, style: str) -> None:
    if OxmlElement is None or qn is None:
        paragraph = document.add_paragraph("目录将在 Word 中更新。")
        _format_body_paragraph(paragraph, style, first_line_indent=False)
        return

    paragraph = document.add_paragraph()
    _format_body_paragraph(paragraph, style, first_line_indent=False)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "右键更新目录"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)

    if hasattr(document, "settings"):
        settings = document.settings.element
        exists = settings.xpath("./w:updateFields")
        if not exists:
            update_fields = OxmlElement("w:updateFields")
            update_fields.set(qn("w:val"), "true")
            settings.append(update_fields)


def _remove_empty_paragraphs(document: Any) -> None:
    """Remove empty paragraphs from document body (post-processing).

    Keeps headings, tables, and paragraphs with actual content.
    Also removes duplicate consecutive empty paragraphs.
    """
    body = document.element.body
    to_remove = []
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "p":
            continue
        # Check if paragraph has any text content
        texts = child.itertext()
        has_content = any(t.strip() for t in texts)
        # Keep paragraphs that contain images/drawings (no text but has inline shape)
        has_drawing = len(child.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")) > 0
        if not has_content and not has_drawing:
            to_remove.append(child)
    for elem in to_remove:
        body.remove(elem)


def _ensure_footer_page_number(section: Any) -> None:
    if OxmlElement is None or qn is None or WD_ALIGN_PARAGRAPH is None:
        return
    footer = section.footer
    # Clear existing paragraphs to avoid duplicates
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    # 参考论文格式: "第 X 页"
    run1 = paragraph.add_run("第 ")
    if Pt is not None:
        run1.font.size = Pt(10.5)
        run1.font.name = "Times New Roman"
        if qn is not None:
            run1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)
    if Pt is not None:
        run.font.size = Pt(10.5)
        run.font.name = "Times New Roman"
        if qn is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run2 = paragraph.add_run(" 页")
    if Pt is not None:
        run2.font.size = Pt(10.5)
        run2.font.name = "Times New Roman"
        if qn is not None:
            run2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _should_insert_page_break(level: int, text: str) -> bool:
    if level not in (1, 2):
        return False
    heading = str(text).strip()
    # Heading 1 chapters always get a page break
    if level == 1:
        return bool(re.match(r"^\d+\.", heading))
    # Heading 2: specific front-matter sections
    return heading in {"摘要", "Abstract", "目录", "参考文献", "致谢", "附录", "Appendix"}


def _is_front_matter_line(text: str) -> bool:
    stripped = str(text).strip()
    return (stripped.startswith("关键词：") or stripped.startswith("Keywords:")
            or stripped.startswith("摘要") or stripped.startswith("Abstract")
            or stripped.startswith("摘  要") or stripped.startswith("ABSTRACT"))


def _presentation_payload(project_root: Path, source: str | Path | None) -> tuple[dict[str, Any], Path]:
    source_path = Path(source).resolve() if source and Path(source).is_absolute() else (
        (project_root / source).resolve() if source else (project_root / PRESENTATION_JSON_PATH).resolve()
    )
    if not source_path.exists():
        raise FileNotFoundError(f"Presentation payload not found: {source_path}")
    return json.loads(source_path.read_text(encoding="utf-8")), source_path


def _safe_slide_text(text: str) -> str:
    text = re.sub(r"\s+", " ", str(text)).strip()
    return xml_escape(text)


def _ppt_text_paragraph(
    text: str,
    *,
    level: int = 0,
    size: int = 2200,
    bold: bool = False,
    bullet: bool = True,
) -> str:
    bullet_xml = '<a:buChar char="&#8226;"/>' if bullet else ""
    return (
        f'<a:p><a:pPr lvl="{level}" marL="{342900 + level * 228600}" indent="-171450">{bullet_xml}</a:pPr>'
        f'<a:r><a:rPr lang="zh-CN" sz="{size}"{" b=\"1\"" if bold else ""}/>'
        f"<a:t>{_safe_slide_text(text)}</a:t></a:r>"
        f'<a:endParaRPr lang="zh-CN" sz="{size}"{" b=\"1\"" if bold else ""}/></a:p>'
    )


def _ppt_shape(shape_id: int, name: str, x: int, y: int, cx: int, cy: int, paragraphs: list[str]) -> str:
    body = "".join(paragraphs) or "<a:p/>"
    return f"""
    <p:sp>
      <p:nvSpPr>
        <p:cNvPr id="{shape_id}" name="{xml_escape(name)}"/>
        <p:cNvSpPr txBox="1"/>
        <p:nvPr/>
      </p:nvSpPr>
      <p:spPr>
        <a:xfrm><a:off x="{x}" y="{y}"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
        <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
        <a:noFill/>
        <a:ln><a:noFill/></a:ln>
      </p:spPr>
      <p:txBody>
        <a:bodyPr wrap="square" rtlCol="0"/>
        <a:lstStyle/>
        {body}
      </p:txBody>
    </p:sp>
    """.strip()


SLIDE_W = 12192000  # 16:9 widescreen (13.333")
SLIDE_H = 6858000   # 7.5"


def _slide_xml(slide: dict, slide_index: int = 0, image_rels: dict[str, str] | None = None) -> str:
    layout = slide.get("layout", "titleAndContent")
    title = str(slide.get("title") or f"Slide {slide_index + 1}")
    bullets = [str(item) for item in (slide.get("bullets") or [])]
    notes = str(slide.get("notes") or "")
    left_col = [str(item) for item in (slide.get("left_column") or [])]
    right_col = [str(item) for item in (slide.get("right_column") or [])]
    w, h = SLIDE_W, SLIDE_H

    layout_fn = _LAYOUT_MAP.get(layout, _layout_title_and_content)
    shapes = layout_fn(title, bullets, notes, left_col, right_col, w, h, slide, image_rels)

    shapes_xml = "\n".join(shapes)
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
       xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
       xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:nvGrpSpPr>
        <p:cNvPr id="1" name=""/>
        <p:cNvGrpSpPr/>
        <p:nvPr/>
      </p:nvGrpSpPr>
      <p:grpSpPr>
        <a:xfrm>
          <a:off x="0" y="0"/>
          <a:ext cx="0" cy="0"/>
          <a:chOff x="0" y="0"/>
          <a:chExt cx="0" cy="0"/>
        </a:xfrm>
      </p:grpSpPr>
      {shapes_xml}
    </p:spTree>
  </p:cSld>
  <p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>
</p:sld>
"""


def _layout_title_only(title, bullets, notes, left_col, right_col, w, h, slide, image_rels):
    title_paras = [_ppt_text_paragraph(title, size=4400, bold=True, bullet=False)]
    shapes = [_ppt_shape(2, "Title", w // 8, h // 4, w * 3 // 4, h // 4, title_paras)]
    if bullets:
        sub_text = "  |  ".join(bullets[:3])
        sub_paras = [_ppt_text_paragraph(sub_text, size=2000, bullet=False)]
        shapes.append(_ppt_shape(3, "Subtitle", w // 8, h // 2 + 200000, w * 3 // 4, 600000, sub_paras))
    if notes:
        shapes.append(_notes_shape(notes, w, h))
    return shapes


def _layout_title_and_content(title, bullets, notes, left_col, right_col, w, h, slide, image_rels):
    title_paras = [_ppt_text_paragraph(title, size=3600, bold=True, bullet=False)]
    shapes = [_ppt_shape(2, "Title", 457200, 200000, w - 914400, 720000, title_paras)]
    body_paras = [_ppt_text_paragraph(item, size=2200) for item in bullets]
    shapes.append(_ppt_shape(3, "Content", 685800, 1050000, w - 1371600, h - 2100000, body_paras))
    if notes:
        shapes.append(_notes_shape(notes, w, h))
    return shapes


def _layout_two_column(title, bullets, left_col, right_col, w, h, slide, image_rels):
    title_paras = [_ppt_text_paragraph(title, size=3600, bold=True, bullet=False)]
    shapes = [_ppt_shape(2, "Title", 457200, 200000, w - 914400, 720000, title_paras)]
    col_w = (w - 914400 - 457200) // 2
    left_paras = [_ppt_text_paragraph(item, size=2000) for item in (left_col or bullets[:3])]
    right_paras = [_ppt_text_paragraph(item, size=2000) for item in (right_col or bullets[3:])]
    shapes.append(_ppt_shape(3, "LeftCol", 457200, 1050000, col_w, h - 2100000, left_paras))
    shapes.append(_ppt_shape(4, "RightCol", 457200 + col_w + 228600, 1050000, col_w, h - 2100000, right_paras))
    notes = str(slide.get("notes") or "")
    if notes:
        shapes.append(_notes_shape(notes, w, h))
    return shapes


def _layout_section_divider(title, bullets, notes, left_col, right_col, w, h, slide, image_rels):
    accent = slide.get("accent_color", "2563EB")
    bg_shape = f"""    <p:sp>
      <p:nvSpPr><p:cNvPr id="2" name="BgRect"/><p:cNvSpPr txBox="1"/><p:nvPr/></p:nvSpPr>
      <p:spPr>
        <a:xfrm><a:off x="0" y="0"/><a:ext cx="{w}" cy="{h}"/></a:xfrm>
        <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
        <a:solidFill><a:srgbClr val="{xml_escape(accent)}"><a:alpha val="15000"/></a:srgbClr></a:solidFill>
        <a:ln><a:noFill/></a:ln>
      </p:spPr>
      <p:txBody><a:bodyPr/><a:lstStyle/><a:p/></p:txBody>
    </p:sp>"""
    title_paras = [_ppt_text_paragraph(title, size=4800, bold=True, bullet=False)]
    shapes = [
        bg_shape,
        _ppt_shape(3, "Title", w // 8, h * 2 // 5, w * 3 // 4, h // 4, title_paras),
    ]
    if bullets:
        sub_text = " — ".join(bullets[:2])
        sub_paras = [_ppt_text_paragraph(sub_text, size=2200, bullet=False)]
        shapes.append(_ppt_shape(4, "Subtitle", w // 8, h * 2 // 5 + h // 4 + 200000, w * 3 // 4, 500000, sub_paras))
    if notes:
        shapes.append(_notes_shape(notes, w, h))
    return shapes


def _layout_image_caption(title, bullets, notes, left_col, right_col, w, h, slide, image_rels):
    title_paras = [_ppt_text_paragraph(title, size=3200, bold=True, bullet=False)]
    shapes = [_ppt_shape(2, "Title", 457200, 150000, w - 914400, 600000, title_paras)]
    img_path = slide.get("image", "")
    if img_path and image_rels and img_path in image_rels:
        r_id = image_rels[img_path]
        ext = Path(img_path).suffix.lstrip(".").lower()
        img_w, img_h = w // 2, h // 2
        img_x = (w - img_w) // 2
        img_y = 900000
        shapes.append(_image_shape(3, "Image", img_x, img_y, img_w, img_h, r_id, ext))
    caption = slide.get("image_caption", "")
    if caption:
        cap_paras = [_ppt_text_paragraph(caption, size=1800, bullet=False)]
        shapes.append(_ppt_shape(4, "Caption", (w - w * 3 // 4) // 2, h // 2 + 1100000, w * 3 // 4, 400000, cap_paras))
    if bullets:
        bullet_paras = [_ppt_text_paragraph(item, size=1800) for item in bullets[:4]]
        shapes.append(_ppt_shape(5, "Bullets", 685800, h * 3 // 4, w - 1371600, h // 4, bullet_paras))
    if notes:
        shapes.append(_notes_shape(notes, w, h))
    return shapes


def _layout_blank(title, bullets, notes, left_col, right_col, w, h, slide, image_rels):
    shapes = []
    if title:
        title_paras = [_ppt_text_paragraph(title, size=2400, bold=False, bullet=False)]
        shapes.append(_ppt_shape(2, "Title", 457200, 200000, w - 914400, 500000, title_paras))
    if notes:
        shapes.append(_notes_shape(notes, w, h))
    return shapes


def _notes_shape(notes: str, w: int, h: int) -> str:
    notes_paras = [_ppt_text_paragraph(notes, size=1400, bullet=False)]
    return _ppt_shape(99, "Notes", 685800, h + 914400, w - 1371600, 914400, notes_paras)


def _image_shape(shape_id: int, name: str, x: int, y: int, cx: int, cy: int, r_id: str, ext: str) -> str:
    return f"""    <p:pic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <p:nvPicPr>
        <p:cNvPr id="{shape_id}" name="{xml_escape(name)}"/>
        <p:cNvPicPr><a:picLocks noChangeAspect="1"/></p:cNvPicPr>
        <p:nvPr/>
      </p:nvPicPr>
      <p:blipFill>
        <a:blip r:embed="{r_id}"/>
        <a:stretch><a:fillRect/></a:stretch>
      </p:blipFill>
      <p:spPr>
        <a:xfrm><a:off x="{x}" y="{y}"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
        <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
      </p:spPr>
    </p:pic>"""


_LAYOUT_MAP = {
    "titleOnly": _layout_title_only,
    "titleAndContent": _layout_title_and_content,
    "twoColumn": _layout_two_column,
    "sectionDivider": _layout_section_divider,
    "imageCaption": _layout_image_caption,
    "blank": _layout_blank,
}


def _slide_rels_xml(extra_rels: list[tuple[str, str, str]] | None = None) -> str:
    rels = ['<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/>']
    if extra_rels:
        for r_id, r_type, target in extra_rels:
            rels.append(f'<Relationship Id="{r_id}" Type="{r_type}" Target="{target}"/>')
    rels_xml = "\n  ".join(rels)
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  {rels_xml}
</Relationships>
"""


def _content_types_xml(slide_count: int, media_exts: set[str] | None = None) -> str:
    defaults = [
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>',
        '<Default Extension="xml" ContentType="application/xml"/>',
    ]
    ext_map = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg"}
    for ext in (media_exts or set()):
        if ext in ext_map:
            defaults.append(f'<Default Extension="{ext}" ContentType="{ext_map[ext]}"/>')
    overrides = [
        '<Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>',
        '<Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/>',
        '<Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/>',
        '<Override PartName="/ppt/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>',
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>',
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>',
    ]
    overrides.extend(
        f'<Override PartName="/ppt/slides/slide{index}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>'
        for index in range(1, slide_count + 1)
    )
    all_items = defaults + overrides
    joined = "\n  ".join(all_items)
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  {joined}
</Types>
"""


def _root_rels_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>
"""


def _core_xml(title: str) -> str:
    created = _now_iso()
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
  xmlns:dc="http://purl.org/dc/elements/1.1/"
  xmlns:dcterms="http://purl.org/dc/terms/"
  xmlns:dcmitype="http://purl.org/dc/dcmitype/"
  xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>{xml_escape(title)}</dc:title>
  <dc:creator>Sci Workspace</dc:creator>
  <cp:lastModifiedBy>Sci Workspace</cp:lastModifiedBy>
  <dcterms:created xsi:type="dcterms:W3CDTF">{created}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{created}</dcterms:modified>
</cp:coreProperties>
"""


def _app_xml(slide_count: int, title: str) -> str:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
  xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>Sci Workspace</Application>
  <PresentationFormat>On-screen Show (16:9)</PresentationFormat>
  <Slides>{slide_count}</Slides>
  <Notes>0</Notes>
  <HiddenSlides>0</HiddenSlides>
  <MMClips>0</MMClips>
  <ScaleCrop>false</ScaleCrop>
  <HeadingPairs>
    <vt:vector size="2" baseType="variant">
      <vt:variant><vt:lpstr>主题</vt:lpstr></vt:variant>
      <vt:variant><vt:i4>1</vt:i4></vt:variant>
    </vt:vector>
  </HeadingPairs>
  <TitlesOfParts>
    <vt:vector size="1" baseType="lpstr">
      <vt:lpstr>{xml_escape(title)}</vt:lpstr>
    </vt:vector>
  </TitlesOfParts>
  <Company>OpenAI</Company>
  <LinksUpToDate>false</LinksUpToDate>
  <SharedDoc>false</SharedDoc>
  <HyperlinksChanged>false</HyperlinksChanged>
  <AppVersion>1.0</AppVersion>
</Properties>
"""


def _presentation_xml(slide_count: int) -> str:
    slide_ids = "\n    ".join(
        f'<p:sldId id="{255 + index}" r:id="rId{index + 1}"/>' for index in range(1, slide_count + 1)
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
  xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:sldMasterIdLst>
    <p:sldMasterId id="2147483648" r:id="rId1"/>
  </p:sldMasterIdLst>
  <p:sldIdLst>
    {slide_ids}
  </p:sldIdLst>
  <p:sldSz cx="12192000" cy="6858000"/>
  <p:notesSz cx="6858000" cy="9144000"/>
</p:presentation>
"""


def _presentation_rels_xml(slide_count: int) -> str:
    slide_rels = "\n  ".join(
        f'<Relationship Id="rId{index + 1}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{index}.xml"/>'
        for index in range(1, slide_count + 1)
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="slideMasters/slideMaster1.xml"/>
  {slide_rels}
</Relationships>
"""


def _slide_master_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldMaster xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
  xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld name="Sci Workspace Theme">
    <p:spTree>
      <p:nvGrpSpPr>
        <p:cNvPr id="1" name=""/>
        <p:cNvGrpSpPr/>
        <p:nvPr/>
      </p:nvGrpSpPr>
      <p:grpSpPr>
        <a:xfrm>
          <a:off x="0" y="0"/>
          <a:ext cx="0" cy="0"/>
          <a:chOff x="0" y="0"/>
          <a:chExt cx="0" cy="0"/>
        </a:xfrm>
      </p:grpSpPr>
    </p:spTree>
  </p:cSld>
  <p:clrMap bg1="lt1" tx1="dk1" bg2="lt2" tx2="dk2" accent1="accent1" accent2="accent2" accent3="accent3" accent4="accent4" accent5="accent5" accent6="accent6" hlink="hlink" folHlink="folHlink"/>
  <p:sldLayoutIdLst>
    <p:sldLayoutId id="1" r:id="rId1"/>
  </p:sldLayoutIdLst>
</p:sldMaster>
"""


def _slide_master_rels_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="../theme/theme1.xml"/>
</Relationships>
"""


def _slide_layout_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldLayout xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
  xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
  type="titleAndContent" preserve="1">
  <p:cSld name="Title and Content">
    <p:spTree>
      <p:nvGrpSpPr>
        <p:cNvPr id="1" name=""/>
        <p:cNvGrpSpPr/>
        <p:nvPr/>
      </p:nvGrpSpPr>
      <p:grpSpPr>
        <a:xfrm>
          <a:off x="0" y="0"/>
          <a:ext cx="0" cy="0"/>
          <a:chOff x="0" y="0"/>
          <a:chExt cx="0" cy="0"/>
        </a:xfrm>
      </p:grpSpPr>
    </p:spTree>
  </p:cSld>
  <p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>
</p:sldLayout>
"""


def _slide_layout_rels_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="../slideMasters/slideMaster1.xml"/>
</Relationships>
"""


def _theme_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Sci Workspace Theme">
  <a:themeElements>
    <a:clrScheme name="Sci Workspace">
      <a:dk1><a:srgbClr val="111827"/></a:dk1>
      <a:lt1><a:srgbClr val="FFFFFF"/></a:lt1>
      <a:dk2><a:srgbClr val="1F2937"/></a:dk2>
      <a:lt2><a:srgbClr val="F8FAFC"/></a:lt2>
      <a:accent1><a:srgbClr val="2563EB"/></a:accent1>
      <a:accent2><a:srgbClr val="0F766E"/></a:accent2>
      <a:accent3><a:srgbClr val="7C3AED"/></a:accent3>
      <a:accent4><a:srgbClr val="EA580C"/></a:accent4>
      <a:accent5><a:srgbClr val="059669"/></a:accent5>
      <a:accent6><a:srgbClr val="DC2626"/></a:accent6>
      <a:hlink><a:srgbClr val="2563EB"/></a:hlink>
      <a:folHlink><a:srgbClr val="7C3AED"/></a:folHlink>
    </a:clrScheme>
    <a:fontScheme name="Sci Workspace">
      <a:majorFont>
        <a:latin typeface="Aptos"/>
        <a:ea typeface="Microsoft YaHei"/>
        <a:cs typeface="Arial"/>
      </a:majorFont>
      <a:minorFont>
        <a:latin typeface="Aptos"/>
        <a:ea typeface="Microsoft YaHei"/>
        <a:cs typeface="Arial"/>
      </a:minorFont>
    </a:fontScheme>
    <a:fmtScheme name="Sci Workspace">
      <a:fillStyleLst>
        <a:solidFill><a:schemeClr val="lt1"/></a:solidFill>
      </a:fillStyleLst>
      <a:lnStyleLst>
        <a:ln w="9525"><a:solidFill><a:schemeClr val="accent1"/></a:solidFill></a:ln>
      </a:lnStyleLst>
      <a:effectStyleLst>
        <a:effectStyle><a:effectLst/></a:effectStyle>
      </a:effectStyleLst>
      <a:bgFillStyleLst>
        <a:solidFill><a:schemeClr val="lt1"/></a:solidFill>
      </a:bgFillStyleLst>
    </a:fmtScheme>
  </a:themeElements>
</a:theme>
"""


def export_presentation_to_pptx(
    project_root: str | Path = ".",
    source: str | Path | None = None,
    output_path: str | Path | None = None,
) -> dict[str, Any]:
    root = Path(project_root).resolve()
    payload, source_path = _presentation_payload(root, source)
    target_path = Path(output_path).resolve() if output_path else _default_output_path(root, source_path, ".pptx")
    slides = payload.get("slides") or []
    if not slides:
        raise ValueError("Presentation payload does not contain slides.")

    # Collect images from slides
    media_files: dict[str, tuple[bytes, str]] = {}  # path -> (data, zip_path)
    slide_image_rels: list[dict[str, str]] = []  # per-slide image rels
    media_exts: set[str] = set()
    img_counter = 0
    for slide in slides:
        img_path_str = slide.get("image", "")
        img_rels: dict[str, str] = {}
        if img_path_str:
            img_path = Path(img_path_str)
            if not img_path.is_absolute():
                img_path = root / img_path
            if img_path.exists():
                ext = img_path.suffix.lstrip(".").lower()
                if ext in ("png", "jpg", "jpeg"):
                    img_counter += 1
                    zip_media_path = f"ppt/media/image{img_counter}.{ext}"
                    media_files[img_path_str] = (img_path.read_bytes(), zip_media_path)
                    r_id = f"rId{2 + len(img_rels)}"
                    img_rels[img_path_str] = r_id
                    media_exts.add(ext)
        slide_image_rels.append(img_rels)

    target_path.parent.mkdir(parents=True, exist_ok=True)
    with ZipFile(target_path, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", _content_types_xml(len(slides), media_exts))
        archive.writestr("_rels/.rels", _root_rels_xml())
        archive.writestr("docProps/core.xml", _core_xml(str(payload.get("title") or "Research Presentation")))
        archive.writestr("docProps/app.xml", _app_xml(len(slides), str(payload.get("title") or "Research Presentation")))
        archive.writestr("ppt/presentation.xml", _presentation_xml(len(slides)))
        archive.writestr("ppt/_rels/presentation.xml.rels", _presentation_rels_xml(len(slides)))
        archive.writestr("ppt/slideMasters/slideMaster1.xml", _slide_master_xml())
        archive.writestr("ppt/slideMasters/_rels/slideMaster1.xml.rels", _slide_master_rels_xml())
        archive.writestr("ppt/slideLayouts/slideLayout1.xml", _slide_layout_xml())
        archive.writestr("ppt/slideLayouts/_rels/slideLayout1.xml.rels", _slide_layout_rels_xml())
        archive.writestr("ppt/theme/theme1.xml", _theme_xml())

        # Write media files
        for img_path_str, (data, zip_path) in media_files.items():
            archive.writestr(zip_path, data)

        for index, slide in enumerate(slides):
            img_rels = slide_image_rels[index] if index < len(slide_image_rels) else {}
            # Build extra rels for images
            extra_rels = []
            for img_path_str, r_id in img_rels.items():
                if img_path_str in media_files:
                    _, zip_path = media_files[img_path_str]
                    extra_rels.append((r_id, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image", f"../{zip_path}"))
            archive.writestr(
                f"ppt/slides/slide{index + 1}.xml",
                _slide_xml(slide, index, img_rels),
            )
            archive.writestr(f"ppt/slides/_rels/slide{index + 1}.xml.rels", _slide_rels_xml(extra_rels))

    export_payload = {
        "kind": "pptx_export",
        "artifact": "presentation",
        "source_path": str(source_path),
        "output_path": str(target_path),
        "slide_count": len(slides),
        "exported_at": _now_iso(),
    }
    json_path = target_path.with_suffix(".pptx.json")
    _write_json(json_path, export_payload)
    export_payload["json_path"] = str(json_path)
    return export_payload


def word_cross_ref_process(input_path: str | Path, output_path: str | Path | None = None) -> dict[str, Any]:
    """Word 交叉引用自动化：将【作者 年份】转为 Word 上标交叉引用 [N]，并按首次出现顺序重排参考文献。"""
    input_path = Path(input_path)
    if output_path is None:
        output_path = input_path.with_stem(input_path.stem + "_cite")
    else:
        output_path = Path(output_path)

    import zipfile as _zf

    with _zf.ZipFile(str(input_path), "r") as z:
        xml = z.read("word/document.xml").decode("utf-8")

    def _get_pure_text(frag: str) -> str:
        text = re.sub(r"<[^>]+>", "", frag)
        return (text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
                .replace("&quot;", '"').replace("&apos;", "'")).strip()

    para_pattern = re.compile(r"<w:p[ >].*?</w:p>", re.DOTALL)
    paras_full = list(para_pattern.finditer(xml))

    # 定位参考文献
    ref_start_idx = -1
    for i, pm in enumerate(paras_full):
        p_text = _get_pure_text(pm.group())
        if re.search(r"^(?:\d+\s*[.．]?\s*)?(参考文献|References)\s*$", p_text, re.I):
            ref_start_idx = i
            break
    if ref_start_idx == -1:
        return {"status": "skipped", "reason": "no_references_section"}

    # 解析参考文献
    parsed_refs: list[dict] = []
    author_map: dict[str, list[dict]] = {}
    bm_id_base = 300000
    new_paras = [p.group() for p in paras_full[:ref_start_idx + 1]]
    seen_refs: set[str] = set()

    for i in range(ref_start_idx + 1, len(paras_full)):
        p_xml = paras_full[i].group()
        p_text = _get_pure_text(p_xml)
        if len(p_text) < 5:
            parsed_refs.append({"valid": False, "xml": p_xml})
            continue
        clean_text = re.sub(r"^(\s*(?:[\[【]\s*\d+\s*[\]】]|\d+\s*[\.．、])\s*)", "", p_text)
        norm_text = re.sub(r"[^\w\u4e00-\u9fff]+", "", clean_text).lower()
        if norm_text:
            if norm_text in seen_refs:
                continue
            seen_refs.add(norm_text)
        author_match = re.search(r"[\u4e00-\u9fff]{2,4}|[a-zA-Z\-]{2,}", clean_text)
        year_match = re.search(r"\b(19|20)\d{2}\b", clean_text)
        if not author_match:
            parsed_refs.append({"valid": False, "xml": p_xml})
            continue
        info = {
            "valid": True, "author": author_match.group().lower(),
            "year": year_match.group() if year_match else None,
            "xml": p_xml, "used": 0, "first_appearance": -1, "bm": "", "bid": 0,
        }
        parsed_refs.append(info)
        author_map.setdefault(info["author"], []).append(info)

    # 替换正文中的【作者 年份】
    match_success = 0
    appearance_counter = 0
    fuzzy_bracket = r"【((?:(?!【|】).)*?)】"

    def _sub_placeholder(m: re.Match) -> str:
        nonlocal match_success, appearance_counter
        content = _get_pure_text(m.group(1)).strip()
        if not content:
            return m.group(0)
        p_author_match = re.search(r"[\u4e00-\u9fff]{2,4}|[a-zA-Z\-]{2,}", content)
        p_year_match = re.search(r"\b(19|20)\d{2}\b", content)
        if not p_author_match:
            return m.group(0)
        p_author = p_author_match.group().lower()
        p_year = p_year_match.group() if p_year_match else None
        target_ref = None
        if p_year and p_author in author_map:
            for r in author_map[p_author]:
                if r["year"] == p_year:
                    target_ref = r
                    break
        if not target_ref and p_author in author_map:
            refs = author_map[p_author]
            min_use = min(r["used"] for r in refs)
            for r in refs:
                if r["used"] == min_use:
                    target_ref = r
                    break
        if not target_ref:
            return m.group(0)
        if target_ref["first_appearance"] == -1:
            appearance_counter += 1
            target_ref["first_appearance"] = appearance_counter
            target_ref["bm"] = f"_AutoRef_{appearance_counter}"
            target_ref["bid"] = bm_id_base + appearance_counter
        target_ref["used"] += 1
        match_success += 1
        n = target_ref["first_appearance"]
        return (
            f'</w:t></w:r><w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
            f'<w:fldChar w:fldCharType="begin"/></w:r><w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
            f'<w:instrText xml:space="preserve"> REF {target_ref["bm"]} \\r \\h </w:instrText></w:r>'
            f'<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:fldChar w:fldCharType="separate"/></w:r>'
            f'<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:t>[{n}]</w:t></w:r>'
            f'<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:fldChar w:fldCharType="end"/></w:r>'
            f'<w:r><w:t xml:space="preserve">'
        )

    final_paras = []
    for i in range(ref_start_idx):
        final_paras.append(re.sub(fuzzy_bracket, _sub_placeholder, new_paras[i], flags=re.DOTALL))
    final_paras.append(new_paras[ref_start_idx])

    # 按首次出现排序
    cited_refs = sorted(
        [r for r in parsed_refs if r.get("valid") and r["first_appearance"] != -1],
        key=lambda x: x["first_appearance"],
    )
    uncited_refs = [r for r in parsed_refs if r.get("valid") and r["first_appearance"] == -1]
    for r in uncited_refs:
        appearance_counter += 1
        r["first_appearance"] = appearance_counter
        r["bm"] = f"_AutoRef_{appearance_counter}"
        r["bid"] = bm_id_base + appearance_counter

    def _update_ref_number(p_xml: str, new_num: int) -> str:
        t_pattern = re.compile(r"(<w:t(?: [^>]*?)?>)(.*?)(</w:t>)", re.DOTALL)
        matches = list(t_pattern.finditer(p_xml))
        full_text = "".join(m.group(2) for m in matches)
        prefix_match = re.search(r"^(\s*(?:[\[【]\s*\d+\s*[\]】]|\d+\s*[\.．、])\s*)", full_text)
        if not prefix_match:
            return p_xml
        chars_to_delete = len(prefix_match.group(1))
        inserted_new = False
        new_xml = ""
        last_end = 0
        for m in matches:
            new_xml += p_xml[last_end:m.start()]
            t_open, t_text, t_close = m.groups()
            if chars_to_delete > 0:
                if len(t_text) <= chars_to_delete:
                    chars_to_delete -= len(t_text)
                    t_text = ""
                else:
                    t_text = t_text[chars_to_delete:]
                    chars_to_delete = 0
            if not inserted_new and chars_to_delete == 0:
                t_text = f"[{new_num}] " + t_text
                inserted_new = True
            new_xml += f"{t_open}{t_text}{t_close}"
            last_end = m.end()
        new_xml += p_xml[last_end:]
        return new_xml

    for r in cited_refs + uncited_refs:
        new_xml = _update_ref_number(r["xml"], r["first_appearance"])
        bm_xml = f'<w:bookmarkStart w:id="{r["bid"]}" w:name="{r["bm"]}"/><w:bookmarkEnd w:id="{r["bid"]}"/>'
        ins = "</w:pPr>" if "<w:pPr>" in new_xml else ">"
        new_xml = new_xml.replace(ins, f"{ins}{bm_xml}", 1)
        final_paras.append(new_xml)

    for r in parsed_refs:
        if not r.get("valid"):
            final_paras.append(r["xml"])

    while len(final_paras) < len(paras_full):
        final_paras.append("")

    # 拼回 XML
    xml_output = xml
    for i in range(len(paras_full) - 1, -1, -1):
        s, e = paras_full[i].span()
        xml_output = xml_output[:s] + final_paras[i] + xml_output[e:]

    # 保存
    _read_path = str(input_path)
    _write_path = str(output_path)
    # When input == output, read everything into memory first to avoid
    # truncating the source on Windows before reading is complete.
    if Path(_read_path).resolve() == Path(_write_path).resolve():
        with _zf.ZipFile(_read_path, "r") as zin:
            _zip_entries = [(item, zin.read(item)) for item in zin.namelist()]
        with _zf.ZipFile(_write_path, "w", _zf.ZIP_DEFLATED) as zout:
            for item, data in _zip_entries:
                if item == "word/document.xml":
                    zout.writestr(item, xml_output.encode("utf-8"))
                else:
                    zout.writestr(item, data)
    else:
        with _zf.ZipFile(_read_path, "r") as zin:
            with _zf.ZipFile(_write_path, "w", _zf.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    if item == "word/document.xml":
                        zout.writestr(item, xml_output.encode("utf-8"))
                    else:
                        zout.writestr(item, zin.read(item))

    leftovers = re.findall(fuzzy_bracket, xml_output, flags=re.DOTALL)
    clean_leftovers = [_get_pure_text(x).strip() for x in leftovers if _get_pure_text(x).strip()]

    return {
        "status": "success",
        "input_path": str(input_path),
        "output_path": str(output_path),
        "citations_converted": match_success,
        "references_total": sum(1 for r in parsed_refs if r.get("valid")),
        "unmatched_markers": clean_leftovers,
    }
